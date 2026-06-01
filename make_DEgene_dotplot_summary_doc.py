#!/usr/bin/env python3
"""
Generate the summary Word document for the Xenium_May2026 DE-gene dot-plot
analysis (dot-plot versions of the cross-cell-type heat maps, V1).

Documents every parameter used and the verbatim prompts that drove the work.
Output path via --outdir. Requires python-docx.
"""
from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def h(doc, text, level):
    return doc.add_heading(text, level=level)


def kv_table(doc, rows, col0='Parameter', col1='Value'):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    hdr[0].paragraphs[0].add_run(col0).bold = True
    hdr[1].paragraphs[0].add_run(col1).bold = True
    for k, v in rows:
        c = t.add_row().cells
        c[0].text = str(k)
        c[1].text = str(v)
    return t


def mono(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    return p


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    return p


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument('--outdir', required=True, type=Path)
    args = ap.parse_args()
    args.outdir.mkdir(parents=True, exist_ok=True)
    today = date.today().isoformat()

    doc = Document()
    # base style
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(10.5)

    title = doc.add_heading('Xenium May 2026 — DE-gene dot plots of the cross-cell-type '
                            'heat maps (V1)', level=0)
    sub = doc.add_paragraph()
    sub.add_run('Analysis summary: parameters & prompts').bold = True
    meta = doc.add_paragraph()
    meta.add_run(f'Naomi Kassahun · MEWS Lab · {today}\n')
    meta.add_run('Project: Xenium_May2026 (mouse hippocampus spatial transcriptomics)')

    # ---------------------------------------------------------------- overview
    h(doc, '1. Overview', 1)
    doc.add_paragraph(
        'Two dot-plot deliverables were produced as dot-plot counterparts of the '
        'cross-cell-type heat maps, for slide-pooled (A+B) Version 1 (broad-ROI) data. '
        'Both reuse the heat maps’ exact anchor gene selection and show those same '
        'genes across three cell types; they differ only in what each dot encodes:')
    doc.add_paragraph(
        'Abundance dot plot — dot size = % of cells positive; dot colour = mean '
        'transcripts per cell (log scale). Columns = the four experimental groups.',
        style='List Bullet')
    doc.add_paragraph(
        'Fold-change dot plot — dot size = % of cells positive (test group); dot '
        'colour = log2 fold change (clipped ±0.6), taken from the same DE table the '
        'heat maps used. Columns = the four contrasts.', style='List Bullet')

    # ---------------------------------------------------------------- dataset
    h(doc, '2. Dataset & data handling', 1)
    kv_table(doc, [
        ('Assay', '10x Xenium spatial transcriptomics, cell-level'),
        ('Tissue', 'Mouse hippocampus'),
        ('Panel', '~5,104-gene panel (15,312 DE rows = 5,104 genes × 3 cell types per contrast)'),
        ('Slides', 'A = 0088372, B = 0088233 (Region_1, 2026-05-13)'),
        ('Version', 'V1 only (broad-ROI annotation punches)'),
        ('Pooling', "ad.concat(join='inner', merge='same'); obs_names prefixed SlideA:: / SlideB::"),
        ('Cell QC', 'filter_cells(min_counts=10); filter_genes(min_cells=5)'),
        ('Raw counts', "layers['counts'] = X.copy() captured BEFORE normalization"),
        ('Normalization', 'normalize_total(target_sum=1e4) then log1p'),
        ('V1 composition (4 groups)', 'Excitatory 44,746 · Astrocyte 26,047 · Inhibitory 2,388 · Other 21,315'),
    ])

    # ---------------------------------------------------------------- celltyping
    h(doc, '3. Cell typing (master build)', 1)
    doc.add_paragraph(
        'Identical to the master pipeline. Marker-module scores via scanpy '
        'score_genes (random_state=0, n_bins=25); coarse cell type = argmax of the '
        'module scores (score ≤ 0 → Unclassified). Within neurons, Excitatory vs '
        'Inhibitory is split by score > 0 with excitatory precedence; astrocytes kept '
        'as their own subtype; everything else → Other.')
    kv_table(doc, [
        ('Neuron markers', 'Rbfox3, Snap25, Syn1, Syt1, Stmn2, Map2, Tubb3'),
        ('Astrocyte markers', 'Gfap, Aqp4, Slc1a3, Aldh1l1, S100b, Aldoc, Gja1'),
        ('Oligodendrocyte markers', 'Mog, Olig1, Olig2, Sox10'),
        ('Microglia markers', 'Cx3cr1, Tmem119, Csf1r, Aif1, Trem2'),
        ('Excitatory markers', 'Slc17a7, Camk2a, Camk2b, Satb2, Tbr1, Neurod6'),
        ('Inhibitory markers', 'Gad1, Gad2, Slc32a1, Pvalb, Sst, Vip, Reln, Lhx6'),
        ('Panels shown', 'Excitatory, Inhibitory, Astrocyte'),
    ])

    # ---------------------------------------------------------------- gene sel
    h(doc, '4. DE-gene (row) selection — identical to the heat maps', 1)
    doc.add_paragraph(
        'Gene rows are chosen from the same long-form DE table the heat maps used '
        '(SlidesAB_volcanos_DE_combined_2026-05-28.csv). For each anchor cell type, in '
        'the EtOH_veh-vs-H2O_veh contrast, genes are kept if |log2FC| ≥ 0.5 AND '
        'adj p < 1e-3, then the top 50 up (descending log2FC) + top 50 down (ascending) '
        'are taken. Those same 100 genes are then shown across all three cell types. '
        'One page per anchor.')
    kv_table(doc, [
        ('Selection contrast', 'EtOH_veh vs H2O_veh'),
        ('Significance', '|log2FC| ≥ 0.5 AND padj < 1e-3'),
        ('Top-N per direction', '50 up + 50 down (100 genes per anchor)'),
        ('Anchors (one page each)', 'Excitatory, Astrocyte'),
        ('Inhibitory anchor', 'omitted (~0 significant genes)'),
        ('Row order', 'strongest up at top → strongest down at bottom; dashed divider'),
    ])

    # ---------------------------------------------------------------- abundance
    h(doc, '5. Deliverable A — Abundance dot plot', 1)
    kv_table(doc, [
        ('Dot SIZE', '% of cells positive for the transcript (counts > 0) in that group × cell type'),
        ('Size mapping', 'area = SIZE_MIN + clip(frac,0,1)×(SIZE_MAX−SIZE_MIN), SIZE_MIN=6, SIZE_MAX=190'),
        ('Dot COLOUR', 'mean transcripts per cell (raw counts, averaged over ALL cells incl. zeros)'),
        ('Colour scale', 'viridis, LogNorm, SHARED across both anchor pages'),
        ('Colour range', 'vmin = max(1st pctile, 0.01) = 0.01; vmax = 99.5th pctile = 7.47'),
        ('Columns', 'H2O_veh (control), H2O_MCT1i, EtOH_veh, EtOH_MCT1i (4 groups)'),
        ('Min-cell guard', 'blank if < 10 cells in that group × cell type'),
        ('Size legend', '% cells positive: 25 / 50 / 75 / 100%'),
    ])

    # ---------------------------------------------------------------- fc
    h(doc, '6. Deliverable B — Fold-change dot plot', 1)
    kv_table(doc, [
        ('Dot SIZE', '% of cells positive (counts > 0) in the TEST group of the contrast × cell type'),
        ('Size mapping', 'same as abundance (SIZE_MIN=6, SIZE_MAX=190)'),
        ('Dot COLOUR', 'log2 fold change, taken from the DE table (identical values to the heat maps)'),
        ('Colour scale', 'RdBu_r diverging, linear Normalize, symmetric'),
        ('Colour clip', '±0.6 (configurable via --clip; full unclipped values kept in the source CSV)'),
        ('Columns (4 contrasts)', 'EtOH vs H2O · MCT1i vs H2O · EtOH+MCT1i vs H2O · EtOH+MCT1i vs EtOH'),
        ('  EtOH vs H2O', 'test=EtOH_veh, ref=H2O_veh'),
        ('  MCT1i vs H2O', 'test=H2O_MCT1i, ref=H2O_veh'),
        ('  EtOH+MCT1i vs H2O', 'test=EtOH_MCT1i, ref=H2O_veh'),
        ('  EtOH+MCT1i vs EtOH', 'test=EtOH_MCT1i, ref=EtOH_veh'),
        ('Min-cell guard', 'blank if < 10 cells in test OR reference group × cell type'),
        ('Size legend', '% cells positive (test group): 25 / 50 / 75 / 100%'),
    ])

    # ---------------------------------------------------------------- figstds
    h(doc, '7. Figure & output standards (MEWS Lab)', 1)
    doc.add_paragraph('Vector PDF; editable text via pdf.fonttype = 42 / ps.fonttype = 42 '
                      '(svg.fonttype = none).', style='List Bullet')
    doc.add_paragraph('No hardcoded paths — all inputs/outputs supplied as CLI arguments.',
                      style='List Bullet')
    doc.add_paragraph('Every figure has a matching source-data CSV (one per anchor page).',
                      style='List Bullet')
    doc.add_paragraph('Cover page + one page per anchor; cell-type column headers repeated '
                      'top and bottom (secondary x-axis) for tall figures.', style='List Bullet')
    doc.add_paragraph('Gene alias annotations: Slc2a1 (GLUT1), Slc2a3 (GLUT3), Cat (catalase).',
                      style='List Bullet')

    # ---------------------------------------------------------------- outputs
    h(doc, '8. Output files', 1)
    doc.add_paragraph('Working directory:')
    mono(doc, '/Users/naomi/Desktop/10x/10 Xenium may 2026 data/analysis/current/')
    doc.add_paragraph('Scripts:')
    for s in ['build_DEgene_dotplot_abundance_V1.py',
              'build_DEgene_dotplot_fc_V1.py',
              'make_DEgene_dotplot_summary_doc.py  (this document)']:
        mono(doc, '  ' + s)
    doc.add_paragraph('Figures & source data (in Dotplots/):')
    for s in [
        'Xenium_DEgeneDotplot_abundance_V1_Summary_2026-05-31.pdf',
        'SlidesAB_DEgeneDotplot_abundance_V1_anchorExcitatory_2026-05-31.csv',
        'SlidesAB_DEgeneDotplot_abundance_V1_anchorAstrocyte_2026-05-31.csv',
        'Xenium_DEgeneDotplot_FC_V1_Summary_2026-05-31.pdf',
        'SlidesAB_DEgeneDotplot_FC_V1_anchorExcitatory_2026-05-31.csv',
        'SlidesAB_DEgeneDotplot_FC_V1_anchorAstrocyte_2026-05-31.csv',
    ]:
        mono(doc, '  Dotplots/' + s)
    doc.add_paragraph('Inputs (read-only):')
    for s in [
        'output-XETG00253__0088372__Region_1__20260513__192251/  (Slide A: cells.csv.gz, cell_feature_matrix.h5)',
        'output-XETG00253__0088233__Region_1__20260513__192252/  (Slide B)',
        'Slide_A_annotations.csv, Slide_B_annotations.csv  (V1 ROI polygons)',
        'SlidesAB_volcanos_DE_combined_2026-05-28.csv  (long-form DE table)',
    ]:
        mono(doc, '  ' + s)

    # ---------------------------------------------------------------- prompts
    h(doc, '9. Prompts given during the analysis', 1)
    doc.add_paragraph('User instructions are reproduced verbatim (original spelling), in order:')
    prompts = [
        ('Dual-encoding curated dot plot (earlier in session)',
         'and then make one where you have both the percent of cells that as the size '
         'of the and then the # of transcipts as the heat'),
        ('Dot-plot versions of the heat maps (initial request)',
         'okay great so for V1 the heat maps you made were for the different. You had '
         'the top 50 genes for exictiory and down for eexitiyor, and then you show it '
         'for all the cell types. inhbitory and asotryctes, could you make a version of '
         'those heat maps but with the dot plots so make the size the percent of cells '
         'and the color the abundance and then use the color max -0.6 and 0.6 and also '
         'if you could at the end of this anaylsis maybe have a summary word document '
         'about all the paramenters used and then also the promted you were given to do '
         'the anaylsis'),
        ('Clarification: abundance first, then fold change; V1 only',
         'so sorry the dot blot is going to be like not the fold change actually just '
         'show the abindance so the groups would be control mct1i _ h20 .... ethanol '
         'and ethanol +MCT1i lets make that verison first than maybe make ones with the '
         'fold change and also now you can make this for only verison one so basically '
         'the point is we want the dot blots of the DE expressed genes in the heat maps'),
        ('Confirmation of fold-change column design (4 contrasts, size=%pos, colour=log2FC ±0.6)',
         'yes thats perfect'),
    ]
    for label, text in prompts:
        p = doc.add_paragraph()
        p.add_run(label + ':').bold = True
        quote(doc, '“' + text + '”')

    doc.add_paragraph(
        'Interpretation notes: the ±0.6 colour maximum applies to the fold-change '
        'dot plot only (the abundance plot uses a log colour scale of mean transcripts). '
        'Columns are the four experimental groups for abundance and the four contrasts '
        'for fold change. Both are V1 only.')

    # ---------------------------------------------------------------- caveats
    h(doc, '10. Caveats', 1)
    doc.add_paragraph(
        'Cell-level Wilcoxon DE: adjusted p-values are inflated by pseudoreplication '
        '(cells within a mouse are not independent); treat direction and log2FC '
        'magnitude as primary. Detection (% positive) and abundance also reflect probe '
        'efficiency — compare the same gene across groups/contrasts, not different '
        'genes against each other. Inhibitory neurons are comparatively sparse '
        '(n=2,388 across the four groups).')

    out_path = args.outdir / f'Xenium_DEgeneDotplot_Analysis_Summary_{today}.docx'
    doc.save(str(out_path))
    print('Wrote', out_path)


if __name__ == '__main__':
    main()
