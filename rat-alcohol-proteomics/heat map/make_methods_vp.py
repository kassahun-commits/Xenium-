from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
    return p

def para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def set_cell(cell, text, size=10, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.bold = bold
    if not cell.paragraphs[0].runs:
        run = cell.paragraphs[0].add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.bold = bold
    else:
        cell.paragraphs[0].runs[0].text = text

# ── Title ──
title = doc.add_heading('Volcano Plot and Heatmap Generation — Methods', 0)
title.runs[0].font.name = 'Arial'
title.runs[0].font.size = Pt(14)

para(doc, 'April 2026', size=10)
doc.add_paragraph()

# ── 1. Overview ──
heading(doc, '1. Overview', level=1)
para(doc,
    'Volcano plots and heatmaps were generated for four subcellular fractionation compartments '
    '(Membrane, Cytosol, Chromatin, Soluble nuclear) across three alcohol exposure conditions '
    '(Intoxication vs Naïve, Acute Withdrawal vs Naïve, Protracted Abstinence vs Naïve). '
    'All analyses were performed in Python using custom scripts. The source data file was: '
    'EDIT  Philipp Alcohol proteome Simplified - Jan 2024  copy.xlsx.')

# ── 2. Protein filtering ──
heading(doc, '2. Protein Filtering', level=1)
para(doc,
    'Prior to statistical analysis, proteins were filtered by subcellular location annotation '
    '(see Subcellular Location Annotation Methods document for full details):')
bullet(doc, 'Chromatin and Soluble nuclear: proteins classified as Keep or Review were retained '
            '(n = 2,192 per compartment). Exclude proteins (membrane-anchored and secreted; '
            'n = 821) were removed as likely fractionation contaminants.')
bullet(doc, 'Membrane and Cytosol: all proteins were retained (n = 3,013) as no contamination '
            'filter was applied to these compartments.')

# ── 3. Statistical analysis ──
heading(doc, '3. Statistical Analysis', level=1)
para(doc,
    'Fold change and corrected p-values were calculated from raw replicate intensity values '
    'for each protein in each condition comparison.')

heading(doc, '3.1 Fold Change', level=2)
para(doc,
    'Fold change was calculated as the difference in mean log2 intensity between the condition '
    'group and the Naïve group:')
para(doc, '    Fold Change = mean(condition replicates) − mean(Naïve replicates)', size=10)
para(doc,
    'Because the raw data values are already on a log2 scale, this difference corresponds to '
    'a log2 fold change.')

heading(doc, '3.2 P-value and Correction', level=2)
para(doc,
    'A two-sample Welch\'s t-test (unequal variance, scipy.stats.ttest_ind with equal_var=False) '
    'was applied to each protein independently, comparing condition replicates to Naïve replicates. '
    'Proteins with fewer than two non-missing values in either group were excluded from testing.')
para(doc,
    'P-values were corrected using a rank-based adjustment replicating the Excel correction formula '
    'used in the original dataset:')
para(doc, '    Corrected = −log2( min(1,  p × n / rank_descending) )', size=10)
para(doc,
    'where n is the total number of proteins with valid p-values, and rank_descending is the '
    'rank of each protein\'s p-value in descending order (largest p = rank 1). This is equivalent '
    'to a Bonferroni-style correction expressed on a −log2 scale.')

# ── 4. Significance thresholds ──
heading(doc, '4. Significance Thresholds', level=1)
para(doc, 'A protein was considered significant if it met both of the following criteria:')
bullet(doc, 'Corrected p-value > 3.3 (−log2 scale)')
bullet(doc, '|Fold Change| > 0.5')
para(doc,
    'Up-regulated proteins: Fold Change > +0.5 and Corrected > 3.3. '
    'Down-regulated proteins: Fold Change < −0.5 and Corrected > 3.3.')

# ── 5. Volcano plots ──
heading(doc, '5. Volcano Plots', level=1)
para(doc,
    'Volcano plots were generated using matplotlib (Python). Each plot displays Fold Change '
    'on the x-axis and Corrected p-value (−log2 adjusted p) on the y-axis.')
bullet(doc, 'Significant up-regulated proteins are shown in coral (#E8735A).')
bullet(doc, 'Significant down-regulated proteins are shown in sage green (#6B9E78).')
bullet(doc, 'Non-significant proteins are shown in light grey (#C8C8C8).')
bullet(doc, 'Dashed threshold lines are drawn at Fold Change = ±0.5 (vertical) and '
            'Corrected p-value = 3.3 (horizontal).')
para(doc,
    'Twelve volcano plots were produced in total (4 compartments × 3 conditions), saved as '
    'a multi-page PDF (Volcano_Plots_Keep_Review.pdf). Each page contains three plots '
    'corresponding to the three conditions for one compartment, in the order: '
    'Membrane, Cytosol, Chromatin, Soluble nuclear.')

# ── 6. Heatmaps ──
heading(doc, '6. Heatmaps', level=1)

heading(doc, '6.1 Protein Selection', level=2)
para(doc,
    'For each compartment, a union set of significant proteins was identified: any protein '
    'classified as Up or Down in at least one of the three conditions was included. '
    'The number of proteins in each union set was:')

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['Compartment', 'Filter applied', 'Proteins in union']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for run in table.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
shade_row(table.rows[0], 'D9E1F2')

data = [
    ('Membrane',        'All proteins',         '387'),
    ('Cytosol',         'All proteins',         '378'),
    ('Chromatin',       'Keep + Review only',   '1,145'),
    ('Soluble nuclear', 'Keep + Review only',   '417'),
]
for i, (comp, filt, n) in enumerate(data, 1):
    row = table.rows[i]
    row.cells[0].text = comp
    row.cells[1].text = filt
    row.cells[2].text = n
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

doc.add_paragraph()

heading(doc, '6.2 Data Preparation', level=2)
para(doc,
    'The raw replicate intensity values for all four condition groups (Naïve, Intoxication, '
    'Acute Withdrawal, Protracted Abstinence) were extracted for each protein in the union set. '
    'Each row (protein) was z-scored across all replicates:')
para(doc, '    z = (value − row mean) / row standard deviation', size=10)
para(doc,
    'Z-scoring was applied so that relative patterns of up- and down-regulation are visible '
    'across conditions regardless of the absolute intensity of each protein. Missing values '
    'were replaced with zero prior to hierarchical clustering only.')

heading(doc, '6.3 Hierarchical Clustering', level=2)
para(doc,
    'Rows (proteins) were clustered using hierarchical clustering with Ward linkage and '
    'Euclidean distance (scipy.cluster.hierarchy.linkage, method="ward"). The resulting '
    'dendrogram leaf order was used to reorder rows so that proteins with similar expression '
    'patterns across conditions are grouped together. Columns (replicates) were not clustered '
    'and remain in their original condition order.')

heading(doc, '6.4 Visualisation', level=2)
para(doc, 'Heatmaps were rendered using matplotlib with the following settings:')
bullet(doc, 'Colormap: custom pastel diverging (soft blue → white → soft rose), '
            'centered at z = 0 and capped at the 98th percentile of absolute z-scores.')
bullet(doc, 'Replicate columns are grouped by condition with pastel-coloured header bands '
            '(blue = Naïve, orange = Intoxication, green = Acute Withdrawal, purple = Protracted Abstinence).')
bullet(doc, 'White vertical lines separate condition groups.')
bullet(doc, 'Figure height scales with the number of proteins so that each protein is '
            'represented by a visible row.')
bullet(doc, 'Gene symbols are displayed as row labels; font size scales inversely with '
            'the number of proteins to avoid overlap.')
para(doc,
    'Output file: Heatmaps_v2.pdf (4 pages, one per compartment, order: '
    'Membrane, Cytosol, Chromatin, Soluble nuclear).')

# ── 7. Software ──
heading(doc, '7. Software and Libraries', level=1)

sw_table = doc.add_table(rows=7, cols=3)
sw_table.style = 'Table Grid'
for i, h in enumerate(['Library', 'Version', 'Purpose']):
    sw_table.rows[0].cells[i].text = h
    for run in sw_table.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
shade_row(sw_table.rows[0], 'D9E1F2')

sw_data = [
    ('Python',     '3.9',  'Scripting language'),
    ('pandas',     '≥1.3', 'Data loading and manipulation'),
    ('numpy',      '≥1.21','Numerical operations'),
    ('scipy',      '≥1.7', 'Welch\'s t-test, hierarchical clustering'),
    ('matplotlib', '≥3.4', 'Volcano plot and heatmap rendering'),
    ('openpyxl',   '≥3.0', 'Excel file output'),
]
for i, (lib, ver, purp) in enumerate(sw_data, 1):
    row = sw_table.rows[i]
    row.cells[0].text = lib
    row.cells[1].text = ver
    row.cells[2].text = purp
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Analysis performed using Claude Code (Anthropic) — April 2026')
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = None
from docx.shared import RGBColor
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

OUT = 'Volcano_Heatmap_Methods.docx'
doc.save(OUT)
print(f'Saved: {OUT}')
