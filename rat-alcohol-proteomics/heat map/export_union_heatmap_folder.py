"""
Exports everything for the Union Heatmap publication folder:
  1. Heatmap_Union_PubQuality.pdf  — the figure
  2. Union_Heatmap_Proteins.xlsx   — one sheet per compartment with all union proteins
  3. Union_Heatmap_Methods.docx    — plain-English methods write-up
"""

import shutil, os
import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE      = os.path.dirname(os.path.abspath(__file__))
OUT_DIR   = os.path.join(BASE, 'Union_Heatmaps_PubQuality')
FILE      = os.path.join(BASE, '..', 'EDIT  Excluding AWM3 Philipp Alcohol proteome Simplified - Jan 2024  copy.xlsx')
HEATMAP   = os.path.join(BASE, 'Heatmap_Union_PubQuality.pdf')

os.makedirs(OUT_DIR, exist_ok=True)

CORR_THRESH = 3.3

COMPARTMENTS = [
    ('Membrane',        'Membrane',        'all'),
    ('Cytosol',         'Cytosol',         'all'),
    ('Chromatin',       'Chromatin',       'keep_review'),
    ('Soluble Nuclear', 'Soluble nuclear', 'all'),
]

CONDITIONS = [
    ('Intoxication',          'Fold change',   'p-value',   'Corrected'),
    ('Acute Withdrawal',      'Fold change.1', 'p-value.1', 'Corrected.1'),
    ('Protracted Abstinence', 'Fold change.2', 'p-value.2', 'Corrected.2'),
]

# ── 1. Copy heatmap PDF ────────────────────────────────────────────────────────
shutil.copy2(HEATMAP, os.path.join(OUT_DIR, 'Heatmap_Union_PubQuality.pdf'))
print('Copied: Heatmap_Union_PubQuality.pdf')

# ── 2. Build Excel protein list ────────────────────────────────────────────────
wb = Workbook()
wb.remove(wb.active)   # remove default blank sheet

# Colour scheme per compartment
COMP_COLORS = {
    'Membrane':        'E8E4F0',
    'Cytosol':         'D6EDDC',
    'Chromatin':       'F5F0D4',
    'Soluble Nuclear': 'D4E8F5',
}

thin = Side(style='thin', color='BBBBBB')
border = Border(left=thin, right=thin, top=thin, bottom=thin)

for disp, sheet, fmode in COMPARTMENTS:
    df = pd.read_excel(FILE, sheet_name=sheet)
    if fmode == 'keep_review' and 'Filter' in df.columns:
        df = df[df['Filter'].isin(['Keep', 'Review'])].reset_index(drop=True)
    else:
        df = df.reset_index(drop=True)

    # Union mask
    union = pd.Series(False, index=df.index)
    for _, fc_col, _, corr_col in CONDITIONS:
        sig = df[corr_col].notna() & (df[corr_col] >= CORR_THRESH)
        union = union | sig
    df_u = df[union].reset_index(drop=True)

    # Sort by AW fold change
    df_u = df_u.sort_values('Fold change.1', ascending=True).reset_index(drop=True)

    ws = wb.create_sheet(title=disp[:31])   # sheet names max 31 chars

    # Header row
    header_fill = PatternFill('solid', fgColor=COMP_COLORS.get(disp, 'EEEEEE'))
    header_font = Font(name='Arial', bold=True, size=10)

    headers = [
        'Accession', 'Gene Symbol', 'Description',
        'Subcellular Location',
        'I FC (Log2)',   'I Corrected p',   'I Sig',
        'AW FC (Log2)',  'AW Corrected p',  'AW Sig',
        'PA FC (Log2)',  'PA Corrected p',  'PA Sig',
        'In Union (any condition)',
    ]
    if fmode == 'keep_review':
        headers.insert(4, 'Filter')

    for col_i, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_i, value=h)
        cell.font     = header_font
        cell.fill     = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border   = border

    # Data rows
    data_font = Font(name='Arial', size=9)
    for row_i, row in df_u.iterrows():
        sig_i  = row['Corrected']   >= CORR_THRESH if pd.notna(row['Corrected'])   else False
        sig_aw = row['Corrected.1'] >= CORR_THRESH if pd.notna(row['Corrected.1']) else False
        sig_pa = row['Corrected.2'] >= CORR_THRESH if pd.notna(row['Corrected.2']) else False

        vals = [
            row.get('Accession', ''),
            row.get('Gene symbol', ''),
            row.get('Description', ''),
            row.get('Subcellular location', ''),
        ]
        if fmode == 'keep_review':
            vals.append(row.get('Filter', ''))
        vals += [
            round(row['Fold change'],   3) if pd.notna(row['Fold change'])   else '',
            round(row['Corrected'],     3) if pd.notna(row['Corrected'])     else '',
            'Yes' if sig_i  else 'No',
            round(row['Fold change.1'], 3) if pd.notna(row['Fold change.1']) else '',
            round(row['Corrected.1'],   3) if pd.notna(row['Corrected.1'])   else '',
            'Yes' if sig_aw else 'No',
            round(row['Fold change.2'], 3) if pd.notna(row['Fold change.2']) else '',
            round(row['Corrected.2'],   3) if pd.notna(row['Corrected.2'])   else '',
            'Yes' if sig_pa else 'No',
            'Yes',
        ]

        excel_row = row_i + 2   # +2 because row 1 = header, row_i is 0-based
        for col_i, val in enumerate(vals, 1):
            cell = ws.cell(row=excel_row, column=col_i, value=val)
            cell.font   = data_font
            cell.border = border
            if val == 'Yes' and headers[col_i - 1].endswith('Sig'):
                cell.fill = PatternFill('solid', fgColor='FFD7D7')

    # Column widths
    col_widths = [14, 12, 40, 22]
    if fmode == 'keep_review':
        col_widths.insert(4, 8)
    col_widths += [12, 14, 7,  12, 14, 7,  12, 14, 7,  20]
    for col_i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(col_i)].width = w

    ws.freeze_panes = 'A2'
    print(f'  Sheet [{disp}]: {len(df_u)} proteins')

excel_path = os.path.join(OUT_DIR, 'Union_Heatmap_Proteins.xlsx')
wb.save(excel_path)
print(f'Saved: Union_Heatmap_Proteins.xlsx')

# ── 3. Write Word methods document ─────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name  = 'Arial'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def add_para(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.style.font.name = 'Arial'
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_bold_line(doc, label, value):
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.name = 'Arial'
    run_b.font.size = Pt(10)
    run_v = p.add_run(value)
    run_v.font.name = 'Arial'
    run_v.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

# Title
title = doc.add_heading('Union Heatmap Panel — Methods & Data Notes', 0)
for run in title.runs:
    run.font.name  = 'Arial'
    run.font.size  = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

doc.add_paragraph()

# Overview
add_heading(doc, '1. Overview', 1)
add_para(doc,
    'This folder contains a publication-quality union heatmap panel showing '
    'Log2 fold-change proteomics data across four subcellular compartments '
    '(Membrane, Cytosol, Chromatin, Soluble Nuclear) and three alcohol '
    'exposure conditions (Intoxication, Acute Withdrawal, Protracted Abstinence).'
)
add_para(doc,
    'Proteins shown are the "union" set: those reaching statistical significance '
    'in at least one of the three conditions. The heatmap color encodes '
    'Log2 fold change vs Naive controls, using a diverging colormap that matches '
    'the volcano plot color scheme (blue = decreased, pink-red = increased).'
)

doc.add_paragraph()

# Data source
add_heading(doc, '2. Data Source', 1)
add_bold_line(doc, 'File: ', 'EDIT  Excluding AWM3 Philipp Alcohol proteome Simplified - Jan 2024  copy.xlsx')
add_para(doc,
    'This version of the dataset excludes the AW-M-3 replicate, which was '
    'identified as a global outlier across all compartments. Fold change and '
    'corrected p-values used in the heatmap are the pre-computed values '
    'already present in the spreadsheet (not recalculated).',
    indent=True
)

doc.add_paragraph()

# Compartment filters
add_heading(doc, '3. Protein Filters Per Compartment', 1)
add_para(doc,
    'Different inclusion criteria were applied per compartment, consistent '
    'with the volcano plot analysis:'
)
rows_table = [
    ('Compartment',   'Filter Applied',                      'Proteins in Dataset'),
    ('Membrane',      'All proteins (no filter)',            '3,013'),
    ('Cytosol',       'All proteins (no filter)',            '3,013'),
    ('Chromatin',     'Keep and Review only (Exclude removed)', '2,192'),
    ('Soluble Nuclear','All proteins (no filter)',           '3,013'),
]
tbl = doc.add_table(rows=len(rows_table), cols=3)
tbl.style = 'Table Grid'
for r_i, row_vals in enumerate(rows_table):
    for c_i, val in enumerate(row_vals):
        cell = tbl.cell(r_i, c_i)
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            if r_i == 0:
                run.bold = True

doc.add_paragraph()

# Significance threshold
add_heading(doc, '4. Significance Threshold (Union Definition)', 1)
add_para(doc,
    'A protein is included in the heatmap if it passes the significance threshold '
    'in at least one of the three conditions (union approach).'
)
add_bold_line(doc, 'Significance criterion: ', 'Corrected p-value (−log2 scale) >= 3.3')
add_para(doc,
    'The corrected p-value is on a −log2 scale (higher = more significant). '
    'A threshold of 3.3 corresponds to a Bonferroni-style rank-corrected p-value. '
    'No fold-change cutoff was applied — all proteins meeting the corrected p '
    'threshold in any condition are shown.',
    indent=True
)

# Counts table
doc.add_paragraph()
add_heading(doc, '5. Protein Counts', 1)
counts = [
    ('Compartment',   'Total Proteins', 'I Sig', 'AW Sig', 'PA Sig', 'Union'),
    ('Membrane',      '3,013',          '277',   '290',    '256',    '651'),
    ('Cytosol',       '3,013',          '346',   '260',    '231',    '680'),
    ('Chromatin',     '2,192 (K+R)',     '530',  '1,076',  '376',   '1,238'),
    ('Soluble Nuclear','3,013',          '276',   '470',    '178',    '770'),
]
tbl2 = doc.add_table(rows=len(counts), cols=6)
tbl2.style = 'Table Grid'
for r_i, row_vals in enumerate(counts):
    for c_i, val in enumerate(row_vals):
        cell = tbl2.cell(r_i, c_i)
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            if r_i == 0:
                run.bold = True
            if c_i == 5 and r_i > 0:
                run.bold = True

doc.add_paragraph()

# Visualization
add_heading(doc, '6. Visualization Details', 1)
items = [
    ('Color scale: ',       'Diverging colormap from blue (#2B7FD4) through white to pink-red (#E8305A), '
                            'matching the volcano plot colors. Clipped at Log2 FC = +/- 3.0.'),
    ('Protein ordering: ',  'Within each compartment, proteins are sorted by Acute Withdrawal '
                            'fold change (ascending left to right).'),
    ('Panel layout: ',      'Four compartments stacked vertically. The width of each band is '
                            'proportional to the number of proteins in that compartment '
                            '(Chromatin is widest at n=1,238; Membrane is narrowest at n=651).'),
    ('Condition rows: ',    'Each compartment band has 3 rows: Intoxication (top), '
                            'Acute Withdrawal (middle), Protracted Abstinence (bottom). '
                            'Rows are separated by thin white lines.'),
    ('Output format: ',     'PDF (vector), 300 DPI. Suitable for journal submission.'),
    ('Script: ',            'make_heatmap_union_pub.py (in the parent heat map folder).'),
]
for label, val in items:
    add_bold_line(doc, label, val)

doc.add_paragraph()

# Files in folder
add_heading(doc, '7. Files in This Folder', 1)
file_items = [
    ('Heatmap_Union_PubQuality.pdf', 'The publication-quality heatmap panel figure.'),
    ('Union_Heatmap_Proteins.xlsx',  'Protein lists: one sheet per compartment, '
                                     'with fold change and corrected p-values for all '
                                     'three conditions. Significant entries highlighted in red.'),
    ('Union_Heatmap_Methods.docx',   'This document.'),
]
for fname, desc in file_items:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(fname + ': ')
    run_b.bold = True
    run_b.font.name = 'Arial'
    run_b.font.size = Pt(10)
    run_v = p.add_run(desc)
    run_v.font.name = 'Arial'
    run_v.font.size = Pt(10)

doc_path = os.path.join(OUT_DIR, 'Union_Heatmap_Methods.docx')
doc.save(doc_path)
print('Saved: Union_Heatmap_Methods.docx')
print(f'\nAll files saved to: {OUT_DIR}')
