from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = 'Arial'
    return p

def add_paragraph(doc, text='', bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    return p

# Title
title = doc.add_heading('Nuclear Localization Evidence — Rat Alcohol Proteomics Proteins', 0)
title.runs[0].font.name = 'Arial'
title.runs[0].font.size = Pt(14)

add_paragraph(doc, 'Proteins identified from the Chromatin fractionation dataset were searched for published evidence of nuclear localization or translocation. Searches were conducted in PubMed and via web queries using terms "[protein name] nuclear localization", "[protein name] nuclear translocation", and "[protein name] nucleus". Human and mouse evidence was accepted. A protein was included if at least one published paper reported nuclear presence.', size=10)

doc.add_paragraph()

# ── SECTION 1: WITH EVIDENCE ──
heading(doc, 'Proteins with Published Nuclear Evidence', level=1)

proteins_with_evidence = [
    ('Ptms', 'Primary nuclear protein. Interacts with linker histone H1 and promotes chromatin decondensation. Nuclear distribution is punctate, excluding dense heterochromatin.', 'Kanninen K et al. (2004) PMID 15716277'),
    ('Pkig', 'Enters the nucleus to bind and inactivate free PKA catalytic subunits, terminating cAMP-induced immediate-early gene expression. Contains a nuclear export signal.', 'Bhatt DL et al. (2004) PMID 15557275'),
    ('Cacnb4', 'Upon neuronal electrical stimulation, the β4 subunit translocates to the nucleus via protein phosphatase 2A interaction and regulates gene transcription. An epilepsy-linked truncation (R482X) abolishes nuclear translocation.', 'Subramanyam P et al. (2013) PMID 23511121'),
    ('Gmfb', 'Immunofluorescence shows primary nuclear expression in nerve cells. Nuclear fraction confirmed by nuclear protein extraction in liver regeneration studies.', 'UniProt P60983; Lai HH et al. (2022) PMC9606832'),
    ('Prkag2', 'The AMPK γ2 subunit translocates to the nucleus under stress conditions. Contains a nuclear localization signal. Nuclear γ2-AMPK suppresses ribosome biogenesis and protects myocardium during ischemia/reperfusion injury.', 'Li J et al. (2017) PMC5659937'),
    ('Dstn', 'Contains a nuclear localization signal; translocates to the nucleus together with actin in response to cellular stress. Member of the ADF/cofilin family.', 'Iida K et al. (1997) PMID 9118250'),
    ('Eif4g2', 'Nuclear in spermatogonia; Importin-13 (truncated isoform tImp13) regulates nuclear export of EIF4G2. When tImp13 is absent, EIF4G2 accumulates in the nucleus.', 'Geles KG et al. (2017) PMID 27993670'),
    ('Skp1', 'Localizes to both nucleus and cytoplasm. Core component of SCF E3 ubiquitin ligase complexes that mediate nuclear ubiquitin-proteasome-dependent proteolysis of cell cycle regulators. Also described as an RNA Pol II elongation factor.', 'Zhang H et al. (2000) PMID 10778750'),
    ('Psmd7', 'The unmodified S12/PSMD7 isoform is found in nuclei of normal cells (modified S12-M variant is predominantly cytosolic). Associates with the nuclear 26S proteasome.', 'Kikuchi J et al. (2004) PMID 15221960'),
    ('Psmd12', 'Subunit of the 19S regulatory lid; present in nuclear proteasomes. The 26S proteasome shows cell cycle-dependent chromatin localization. Nuclear PSMD12 expression is elevated in cancer cells.', 'Choo MK et al. (2020) PMID 32348834'),
    ('Ugdh', 'Nuclear UGDH expression is observed in lung adenocarcinoma cells and correlates with reduced differentiation and advanced stage. Immunohistochemistry confirms nuclear staining in clinical samples.', 'Shao C et al. (2019) PMID 30787260'),
    ('Phgdh', 'Nuclear PHGDH has a non-enzymatic oncogenic function: its ACT domain binds nuclear cMyc to form a transcriptional complex (PHGDH/p300/cMyc/AF9) that drives CXCL1/IL-8 expression and reshapes the tumor immune microenvironment in liver cancer.', 'Zhu J et al. (2023) PMID 37078828'),
    ('Tkt', 'Found in both cytoplasm and nucleus of hepatocellular carcinoma cells. Nuclear TKT interacts with STAT1 and recruits HDAC3 to the FXR promoter, suppressing FXR transcription.', 'He F et al. (2020) PMID 31949131'),
    ('Mat1a', 'Nuclear accumulation of MATα1 (Mat1a protein) observed in damaged/diseased hepatocytes. Nuclear MATα1 interacts with the oncogene PDRG1, confirmed by subcellular fractionation and immunofluorescence.', 'Peng H et al. (2016) PMID 27548429'),
    ('Fbp1', 'Under glucose deprivation, PERK-mediated phosphorylation of FBP1 at S170 exposes a nuclear localization signal and drives nuclear translocation. Nuclear FBP1 binds PPARα and acts as a histone H3 phosphatase.', 'Zhao Y et al. (2022) PMID 36266488'),
    ('Khk', 'Immunohistochemistry in normal liver hepatocytes and renal tubule cells shows both cytoplasmic and nuclear staining for the KHK-C isoform.', 'Diggle CP et al. (2009) PMID 19365088'),
    ('Hspd1', 'Nuclear HSP60 reported to bind regulatory regions of Oct4, Nanog, c-Myc, p53, TERT, and STAT3 in stem cells. Note: primary paper (PMID 21995449) was retracted; nuclear HSP60 presence is corroborated in independent reviews.', 'Review: PMC5920047 (include with caution)'),
    ('Tpd52', 'Nuclear localization of TPD52 observed in mucinous and clear cell ovarian carcinoma subtypes; may have subtype-specific oncogenic roles in the nucleus.', 'Byrne JA et al. (2005) PMID 15986428'),
    ('Dhrs4', 'A splice isoform (NRDRA2) lacks the peroxisomal targeting signal and gains a C-terminal nuclear localization signal. GFP reporter assays confirm predominantly nuclear localization of NRDRA2.', 'Wang YQ et al. (2012) PMID 22227495'),
    ('Cat', 'Catalase detected in the nuclear matrix of guinea pig hepatocytes. Intrinsic nuclear targeting confirmed using split-fluorescent protein technology; C-terminal tagging was shown to interfere with native nuclear localization.', 'Nguyen MK et al. (1988) PMID 3396586; Costello JL et al. (2022) PMID 35040158'),
    ('Kifap3', 'Found in the nuclear fraction. Forms a ternary complex with chromosome-associated protein HCAP and KIF3B motor protein at chromosomes.', 'Yamazaki H et al. (1998) PMID 9506951'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Protein'
hdr[1].text = 'Evidence'
hdr[2].text = 'Reference'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'D9E1F2')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

col_widths = [Inches(1.0), Inches(3.6), Inches(2.0)]
for i, width in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = width

for protein, evidence, ref in proteins_with_evidence:
    row = table.add_row().cells
    row[0].text = protein
    row[1].text = evidence
    row[2].text = ref
    for j, cell in enumerate(row):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)

doc.add_paragraph()

# ── SECTION 2: NO EVIDENCE ──
heading(doc, 'Proteins with No Nuclear Evidence Found', level=1)
add_paragraph(doc, 'The following proteins were searched and no published evidence of nuclear localization or translocation was identified. This does not exclude nuclear presence — it reflects the limits of the literature search.', size=10)
doc.add_paragraph()

no_evidence = [
    'Grsf1', 'Sos1', 'Ube2v2', 'Wdr7', 'Acsf3', 'Dcakd', 'Akr1c13', 'Hdhd2', 'Mrpl2',
    'Tagln3', 'Map3k9', 'Fn3k', 'Rasgrf1', 'Tufm', 'Idh3a', 'Glb1', 'Lactb2', 'Mybpc3',
    'Pptc7', 'Fam131b', 'Cobll1', 'L2hgdh', 'Dbt', 'Ralgapb', 'Gdap1l1', 'Pde4c', 'Ttc39c',
    'Relch', 'Acox2', 'Vsnl1', 'Bphl', 'Cgn', 'Gart', 'Mrpl15', 'Aass', 'Bpnt1', 'Ube2d2',
    'Pgp', 'Ppm1f', 'Thumpd1', 'Hibch', 'Pah', 'Acat1', 'Acot4', 'Krt79', 'Glyat', 'Grhpr',
    'Ina', 'Rab3il1', 'Hgd', 'Pnpo', 'Srr', 'Acaa1a', 'Qdpr', 'Aldh4a1', 'Acsm5', 'Suclg1',
    'Suclg2', 'Pc', 'Ehhadh', 'Acadsb', 'Asl', 'Otc', 'Phyhd1', 'Hmgcs2', 'Coq9', 'Ech1',
    'Uroc1', 'Hmgcl', 'Etfb', 'Phyh', 'Bckdha', 'Gcdh', 'Ak3', 'Hacl1', 'Decr1', 'Qprt',
    'Myl12b', 'Dmgdh', 'Agxt', 'Iars2', 'Csad', 'Gss', 'Nadk2', 'Etfa', 'Ivd', 'Mccc2',
    'Aldh2', 'Ces2a', 'Sardh', 'Prps1', 'Acad8', 'Iah1', 'Acsf2', 'Prodh', 'Acadm', 'Hibadh',
    'Tdo2', 'Amdhd1', 'Pcca', 'Mrps26', 'Oxsm', 'Acaa1b', 'Agxt2', 'Aadat', 'Eci2', 'Hsd17b4',
    'Adhfe1', 'Akr1c6', 'Acads', 'Tst', 'Acadl', 'Aldh1b1', 'Fah', 'Aldh6a1', 'Pccb', 'Shc3',
    'Hpcal4', 'Rplp2', 'Pgm3',
]

# Format as a wrapped comma-separated paragraph
p = doc.add_paragraph(', '.join(no_evidence))
p.runs[0].font.name = 'Arial'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ── SECTION 3: NOTES ──
heading(doc, 'Notes on Methodology', level=1)
notes = [
    'Searches were performed using PubMed and web queries with terms "[gene name] nuclear localization", "[gene name] nuclear translocation", and "[gene name] nucleus".',
    'Human and mouse literature was accepted as evidence; rat-specific studies were preferred where available.',
    'A protein was included if at least one peer-reviewed publication reported nuclear localization, nuclear translocation, or a nuclear function.',
    'Hspd1 (HSP60): the primary paper reporting nuclear binding of stemness transcription factors (PMID 21995449) was retracted; this entry should be treated with caution and verified independently.',
    'Ufc1: the ufmylation pathway has documented nuclear roles via its E3 partner UFL1, but nuclear localization of UFC1 itself was not directly demonstrated and was excluded from the main list.',
    'Pdk1: the pyruvate dehydrogenase complex (PDC) translocates to the nucleus (PDHA1 specifically), but PDK1 kinase itself was not shown to be nuclear and was excluded.',
]
for note in notes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(note)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

OUT = 'Nuclear_Localization_Evidence.docx'
doc.save(OUT)
print(f'Saved: {OUT}')
