from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
    return p

def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def set_cell_font(cell, size=9, bold=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.bold = bold

# ── Title ──
title = doc.add_heading('Nuclear Localization Evidence — Mitochondrial/Metabolic Protein List', 0)
title.runs[0].font.name = 'Arial'
title.runs[0].font.size = Pt(14)

p = doc.add_paragraph()
run = p.add_run(
    'Proteins from the rat alcohol proteomics Chromatin fractionation dataset were searched for published '
    'evidence of nuclear localization or translocation. Searches used PubMed and web queries with terms '
    '"[gene name] nuclear localization", "[gene name] nuclear translocation", and "[gene name] nucleus". '
    'Human and mouse evidence was accepted. A protein was included if at least one peer-reviewed publication '
    'reported nuclear presence or translocation.'
)
run.font.name = 'Arial'
run.font.size = Pt(10)

doc.add_paragraph()

# ── Section 1: Evidence found ──
set_heading(doc, 'Proteins with Published Nuclear Evidence', level=1)

proteins = [
    ('Grsf1', 'GRSF1 isoform 2 lacks the mitochondrial localization sequence and localizes diffusely to the cytosol and nucleus. GRSF1 also functions in nuclear mRNA export during viral infection.', 'Jourdain AA et al. (2013) PMID 23473034; PMID 39066299', 'Strong'),
    ('Aldh2', 'AMPK/AICAR-induced nuclear translocation demonstrated in rat kidney tissue by subcellular fractionation. Nuclear ALDH2 associates with histone deacetylases and acts as a transcriptional repressor.', 'Choi SE et al. (2011) PMID 21130747', 'Strong'),
    ('Hspe1', 'Nuclear HSP10 (HSPE1) detected by immunofluorescence in bronchial mucosa cells. Nuclear levels increase after cigarette smoke extract stimulation, indicating cytosol-to-nucleus migration.', 'Corrao S et al. (2014) PMID 25355063', 'Strong'),
    ('Acat1', 'Under immune stimulation, mitochondrial ACAT1 is phosphorylated at Ser60 and translocates to the nucleus, where it acetylates p50 (NFKB1) at Lys146, promoting antitumor NK cell immunity.', 'Wei et al. (2025) PMID 40289129', 'Strong'),
    ('Acaa2', 'ACAA2 localizes to the nucleus and acts as a ligand-dependent coactivator for thyroid hormone receptor TRβ1 in cardiac tissue. Confirmed by GST pull-down and luciferase reporter assay in mouse heart.', 'Wang W & Ledee D (2021) PMID 34474245', 'Strong'),
    ('Gcdh', 'In glioblastoma stem cells, GCDH translocates to the nucleus where it interacts with the crotonyltransferase CBP (CREBBP) to promote histone H4 lysine crotonylation (H4Kcr), reprogramming tumor immunity.', 'Yuan H et al. (2023) PMID 37198486 — Nature', 'Strong'),
    ('Mrpl2', 'Dual localization in non-small cell lung cancer: nuclear MRPL2 interacts with PDCD11 to regulate intracellular calcium signaling and NSCLC progression.', 'Jin X et al. (2026) PMID 41558196', 'Strong'),
    ('Decr1', 'Nucleoplasm listed as a confirmed subcellular location in UniProt (Q16698). Categorized as nuclear in single-nucleus RNA-seq subcellular analysis.', 'UniProt Q16698; PMID 41422136', 'Moderate'),
    ('Mrpl15', 'Subcellular distribution analysis (Human Protein Atlas-based bioinformatics) categorizes MRPL15 as predominantly nuclear in gestational diabetes mellitus data.', 'Zhao X et al. (2025) PMID 41334450', 'Weak — bioinformatics only'),
    ('Hspd1', 'Nuclear HSP60 reported to bind Oct4, Nanog, c-Myc, p53, TERT, and STAT3 in stem cells. Note: primary paper (PMID 21995449) was retracted; nuclear HSP60 is corroborated in independent reviews.', 'Review: PMC5920047 (include with caution — primary paper retracted)', 'Weak — retracted primary paper'),
    ('Dhrs4', 'A splice isoform (NRDRA2) lacks the peroxisomal targeting signal and gains a C-terminal nuclear localization signal. GFP reporter assays confirm predominantly nuclear localization of NRDRA2.', 'Wang YQ et al. (2012) PMID 22227495', 'Strong'),
    ('Cat', 'Catalase detected in the nuclear matrix of guinea pig hepatocytes in addition to peroxisomes. Intrinsic nuclear targeting confirmed using split-fluorescent protein technology; C-terminal tagging interferes with native nuclear localization.', 'Nguyen MK (1988) PMID 3396586; Costello JL et al. (2022) PMID 35040158', 'Strong'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, h in enumerate(['Protein', 'Nuclear Evidence', 'Reference', 'Strength of Evidence']):
    hdr_cells[i].text = h
    set_cell_font(hdr_cells[i], size=9, bold=True)
shade_row(table.rows[0], 'D9E1F2')

for protein, evidence, ref, strength in proteins:
    row = table.add_row()
    row.cells[0].text = protein
    row.cells[1].text = evidence
    row.cells[2].text = ref
    row.cells[3].text = strength
    for cell in row.cells:
        set_cell_font(cell, size=9)
    # Color by strength
    if 'Strong' in strength:
        shade_row(row, 'E2EFDA')   # light green
    elif 'Moderate' in strength:
        shade_row(row, 'FFF2CC')   # light yellow
    else:
        shade_row(row, 'FCE4D6')   # light orange/red

col_widths = [Inches(0.8), Inches(3.2), Inches(2.0), Inches(1.1)]
for i, w in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = w

doc.add_paragraph()

# ── Section 2: No evidence ──
set_heading(doc, 'Proteins with No Nuclear Evidence Found', level=1)
note = doc.add_paragraph()
note.add_run(
    'The following proteins were searched and no published evidence of nuclear localization '
    'or translocation was identified. This does not exclude nuclear presence — it reflects the '
    'limits of current literature.'
).font.name = 'Arial'
note.runs[0].font.size = Pt(10)

doc.add_paragraph()

no_evidence = [
    'Acsf3', 'Tufm', 'Idh3a', 'Glb1', 'Lactb2', 'Pptc7', 'L2hgdh', 'Dbt', 'Pdk1',
    'Pde4c', 'Relch', 'Acox2', 'Bphl', 'Cgn', 'Aass', 'Hibch', 'Acot4', 'Glyat',
    'Acaa1a', 'Acaa1b', 'Aldh4a1', 'Acsm5', 'Suclg1', 'Mecr', 'Pccb', 'Txn2', 'Dmgdh',
    'Agxt', 'Iars2', 'Nadk2', 'Etfa', 'Etfb', 'Ivd', 'Mccc2', 'Ces2a', 'Sardh',
    'Bckdha', 'Ak3', 'Hacl1', 'Acsm1', 'Acad8', 'Acsf2', 'Prodh', 'Acadm', 'Hibadh',
    'Pcca', 'Mrps26', 'Oxsm', 'Agxt2', 'Aadat', 'Eci2', 'Hsd17b4', 'Adhfe1', 'Acads',
    'Tst', 'Acadl', 'Aldh1b1', 'Suclg2', 'Pc', 'Ehhadh', 'Acadsb', 'Otc', 'Hmgcs2',
    'Aldh6a1', 'Coq9', 'Ech1', 'Hmgcl', 'Phyh',
]

p = doc.add_paragraph(', '.join(no_evidence))
p.runs[0].font.name = 'Arial'
p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ── Section 3: Notes ──
set_heading(doc, 'Notes on Methodology and Evidence Grading', level=1)
notes = [
    'Searches were performed using PubMed and web queries: "[gene name] nuclear localization", "[gene name] nuclear translocation", "[gene name] nucleus". Human and mouse evidence was accepted.',
    'Strong evidence: direct experimental demonstration (subcellular fractionation, immunofluorescence, co-IP in nuclear fraction, or GFP reporter assay with NLS).',
    'Moderate evidence: confirmed annotation in UniProt with supporting omics data.',
    'Weak evidence: bioinformatics/database annotation only, or primary experimental paper was retracted.',
    'Hspd1 (HSP60): the primary nuclear paper (PMID 21995449) was retracted — verify independently before citing.',
    'Pdk1: the pyruvate dehydrogenase complex (PDC) translocates to the nucleus (PDHA1 subunit specifically), but PDK1 kinase itself was not shown to be nuclear.',
    'Several proteins (e.g., Bckdha, Hmgcl, Acaa1a) regulate nuclear signaling pathways indirectly via metabolite production (acetyl-CoA, β-hydroxybutyrate) without themselves entering the nucleus.',
]
for note in notes:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(note)
    run.font.name = 'Arial'
    run.font.size = Pt(10)

OUT = 'Nuclear_Localization_Evidence_List2.docx'
doc.save(OUT)
print(f'Saved: {OUT}')
