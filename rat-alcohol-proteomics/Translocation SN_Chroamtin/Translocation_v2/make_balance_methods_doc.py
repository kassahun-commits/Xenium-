from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

OUT = 'Chrom_vs_SN_Balance_Methods.docx'

df = pd.read_excel('Chrom_vs_SN_Balance.xlsx', sheet_name='Chrom_vs_SN_Balance')
both_sig = df[(df['Sig_Naive'] != 'NS') & (df['Sig_AW'] != 'NS')]
switchers = both_sig[
    ((both_sig['Sig_Naive'] == 'Up')   & (both_sig['Sig_AW'] == 'Down')) |
    ((both_sig['Sig_Naive'] == 'Down') & (both_sig['Sig_AW'] == 'Up'))
]
n_total  = len(df)
n_both   = len(both_sig)
n_sw     = len(switchers)
n_sn_to_chrom = ((switchers['Sig_Naive']=='Up') & (switchers['Sig_AW']=='Down')).sum()
n_chrom_to_sn = ((switchers['Sig_Naive']=='Down') & (switchers['Sig_AW']=='Up')).sum()
n_aw_up  = (df['Sig_AW']=='Up').sum()
n_aw_dn  = (df['Sig_AW']=='Down').sum()

doc = Document()

# Fonts
def style_run(run, size=11, bold=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14 if level == 1 else 12)
    return p

def add_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, size=size)
    p.paragraph_format.space_after = Pt(8)
    return p

# Title
t = doc.add_heading('Chromatin vs. Soluble Nuclear Compartment Balance Analysis', level=1)
for run in t.runs:
    run.font.name = 'Arial'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Section 1
add_heading(doc, '1. Rationale and Conceptual Framework', level=1)
add_para(doc,
    'A key question in the study of alcohol-related neuroadaptations is whether proteins physically '
    'redistribute between nuclear subcompartments during acute withdrawal (AW). While traditional '
    'approaches compare each fraction against a naïve baseline independently, this analysis takes '
    'a complementary approach: for each protein and each condition, we directly quantify the '
    'relative enrichment of that protein in the chromatin fraction versus the soluble nuclear (SN) '
    'fraction. This "compartment balance" score captures not just whether a protein changes in '
    'abundance overall, but whether it is disproportionately concentrated in one compartment '
    'relative to the other — which is the defining feature of nuclear translocation.')

# Section 2
add_heading(doc, '2. Data Source and Preprocessing', level=1)
add_para(doc,
    'Protein abundance values (log2 LFQ intensities) were drawn from two sheets of the master '
    'proteomics dataset: the Chromatin fraction and the Soluble Nuclear fraction. The Chromatin '
    'sheet was filtered to retain only proteins designated "Keep" or "Review" in the Filter '
    'column, consistent with all prior analyses. The Soluble Nuclear sheet was used without '
    'additional filtering. Animal AW-M-3 was excluded from the Acute Withdrawal group as a '
    'previously identified global outlier exhibiting systematic signal depletion across all '
    'fractions. The final analysis included all proteins detected in both the Chromatin '
    f'(Keep/Review) and Soluble Nuclear sheets, yielding {n_total:,} proteins.')

# Section 3
add_heading(doc, '3. Compartment Balance Score and Statistical Testing', level=1)
add_para(doc,
    'For each protein and each condition (Naïve, Intoxication, Acute Withdrawal, Protracted '
    'Abstinence), a compartment balance fold change was computed as:')
add_para(doc,
    '        FC = mean(Chromatin reps) − mean(Soluble Nuclear reps)')
add_para(doc,
    'where both values are log2 LFQ intensities. Since the data are already on a log2 scale, this '
    'difference is equivalent to log2(Chromatin / SN), i.e., the log2 ratio of chromatin '
    'to soluble nuclear abundance for that condition. A positive FC indicates the protein is more '
    'concentrated in chromatin than in the soluble nuclear fraction; a negative FC indicates the '
    'reverse.')
add_para(doc,
    'Statistical significance was assessed using a two-sample Welch\'s t-test '
    '(scipy.stats.ttest_ind, equal_var=False) comparing the chromatin replicate values against '
    'the soluble nuclear replicate values within each condition. The Acute Withdrawal group '
    'comprised four replicates (AW-F-1, AW-F-2, AW-M-1, AW-M-2); all other conditions comprised '
    'five replicates. P-values were corrected for multiple comparisons using the '
    'Benjamini-Hochberg (BH) procedure, implemented identically to all prior analyses in this '
    'dataset: p-values were ranked in descending order, and the corrected statistic was computed '
    'as −log2(min(1, p × N / rank)), where N is the number of valid tests. A protein was '
    'considered significantly enriched in chromatin (Up) or soluble nuclear (Down) if the '
    'corrected statistic exceeded 3.3 and the absolute fold change exceeded 0.5 log2 units.')

# Section 4
add_heading(doc, '4. Results', level=1)
add_para(doc,
    f'Of the {n_total:,} proteins quantified in both fractions, {n_aw_up} showed significant '
    f'enrichment in chromatin over soluble nuclear during acute withdrawal (AW Up), and {n_aw_dn} '
    f'showed significant enrichment in the soluble nuclear fraction (AW Down).')
add_para(doc,
    f'To identify proteins where the chromatin/SN balance was significant in both the naïve state '
    f'and during acute withdrawal, the two groups were intersected, yielding {n_both} proteins '
    f'significant under both conditions. These proteins represent the set where compartment '
    f'partitioning is robust and detectable in baseline animals as well as during withdrawal.')
add_para(doc,
    f'Of particular interest are proteins that switch their predominant compartment between the '
    f'naïve state and acute withdrawal — that is, proteins that are significantly enriched in one '
    f'fraction under naïve conditions but shift to the opposite fraction during AW. A total of '
    f'{n_sw} proteins showed this pattern: {n_sn_to_chrom} proteins were enriched in chromatin '
    f'over SN in the naïve state but became more SN-enriched during AW, and {n_chrom_to_sn} '
    f'proteins showed the reverse pattern (SN-enriched in naïve, chromatin-enriched in AW). '
    f'These direction-switching proteins represent the strongest candidates for AW-induced '
    f'nuclear translocation.')

# Section 5
add_heading(doc, '5. Heatmap Visualization', level=1)
add_para(doc,
    'Results were visualized as heatmaps displaying the compartment balance fold change '
    '(Chrom − SN) in two columns: Naïve and Acute Withdrawal. Three heatmap versions were '
    'generated:')
add_para(doc,
    '(1) All 683 proteins significant in both Naïve and AW, hierarchically clustered by Ward '
    'linkage on Euclidean distance. This view reveals the overall structure of compartment '
    'partitioning across conditions.')
add_para(doc,
    '(2) The top 50 proteins ranked by absolute AW fold change, sorted from most chromatin-'
    'enriched to most SN-enriched. Gene labels are fully readable at this scale.')
add_para(doc,
    f'(3) The {n_sw} direction-switching proteins, sorted by group (naïve Chrom-enriched → AW '
    f'SN-enriched, then naïve SN-enriched → AW Chrom-enriched) and by AW fold change within '
    f'each group. A dashed white line separates the two switching classes.')
add_para(doc,
    'In all heatmaps, red indicates chromatin enrichment (positive FC) and blue indicates '
    'soluble nuclear enrichment (negative FC), using a diverging colormap centered at zero.')

# Section 6
add_heading(doc, '6. Interpretation', level=1)
add_para(doc,
    'This analysis provides a direct, within-condition view of how proteins partition between '
    'chromatin and the soluble nuclear compartment, and how that partitioning changes during '
    'acute withdrawal from alcohol. Unlike approaches that test each fraction independently '
    'against naïve, the compartment balance score is a single unified metric that captures '
    'the relative distribution of a protein across both nuclear subcompartments simultaneously. '
    'Proteins that switch their predominant compartment specifically during AW — but not under '
    'naïve conditions — are the strongest candidates for condition-specific nuclear translocation '
    'and may represent functionally important chromatin-regulatory events that accompany '
    'withdrawal-induced neuroadaptation.')

doc.save(OUT)
print(f'Saved: {OUT}')
