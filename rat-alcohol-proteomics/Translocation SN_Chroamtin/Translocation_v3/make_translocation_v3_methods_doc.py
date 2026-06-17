"""
Generate a Word document explaining the Translocation v3 analysis,
statistical methods, and heatmap interpretation.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DOCX   = os.path.join(SCRIPT_DIR, 'Translocation_v3_Methods_and_Interpretation.docx')

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=(26, 58, 92))   # dark navy
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3A5C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(44, 82, 130))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, italic=True, size=11, color=(80, 80, 80))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    return p

def mixed(parts):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=11)
    return p

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Protein Translocation into Chromatin During Acute Alcohol Withdrawal')
set_font(run, bold=True, size=16, color=(26, 58, 92))
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Analysis Methods, Statistical Framework, and Heatmap Interpretation')
set_font(run, italic=True, size=12, color=(100, 100, 100))
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Translocation v3  |  MEWS Lab')
set_font(run, size=10, color=(130, 130, 130))
p.paragraph_format.space_after = Pt(18)

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading1('1.  Overview and Scientific Question')

body(
    'This analysis asks a specific question: during Acute Withdrawal (AW) from chronic '
    'alcohol exposure, do certain proteins move from the Soluble Nuclear (SN) compartment '
    'into the Chromatin-bound compartment more than would be expected by a general change '
    'in protein abundance?'
)
body(
    'The key distinction is that a protein could simply be more abundant in the cell overall '
    'during withdrawal — which would make it appear higher in both the SN and Chromatin '
    'fractions simultaneously without any true relocalization. This analysis is specifically '
    'designed to separate true translocation (a protein physically moving from one '
    'compartment to the other) from a general abundance change.'
)
body(
    'To do this, we measure how much each protein changes in Chromatin relative to naive '
    'animals, and compare that change directly to how much the same protein changes in SN '
    'relative to naive — using the AW replicates as the unit of comparison. A protein that '
    'truly translocates into Chromatin will show a large increase in Chromatin and a '
    'smaller increase (or a decrease) in SN, not simply increases in both.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. DATA AND PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════
heading1('2.  Data Source and Preprocessing')

heading2('2.1  Source File')
body(
    'EDIT  Excluding AWM3 Philipp Alcohol proteome Simplified - Jan 2024  copy.xlsx'
)
bullet('Sheet: Chromatin  —  proteins from the chromatin-bound nuclear fraction')
bullet('Sheet: Soluble nuclear  —  proteins from the soluble nuclear fraction')

heading2('2.2  Sample Exclusion')
body(
    'Animal AW-M-3 was excluded from all analyses because it was absent from the source '
    'file (already excluded prior to this analysis due to a quality control issue identified '
    'by the proteomics core).'
)

heading2('2.3  Protein Filtering')
bullet('Chromatin fraction: only proteins with Filter = "Keep" or "Review" were included. '
       'Proteins flagged "Exclude" were removed.')
bullet('Soluble Nuclear fraction: no filter was applied — all detected proteins were included.')
bullet('Only proteins present in both fractions were tested (intersection), yielding 2,190 '
       'proteins with sufficient data for the statistical test.')

heading2('2.4  Replicates Used')

body('Acute Withdrawal (AW) replicates — used for the main statistical test:')
bullet('Chromatin AW: AW-F-1, AW-F-2, AW-M-1, AW-M-2  (4 animals)')
bullet('Soluble Nuclear AW: AW-F-1, AW-F-2, AW-M-1, AW-M-2  (4 animals)')

body('Naive replicates — used to establish baseline:')
bullet('Chromatin Naive: N-F-1, N-F-2, N-F-3, N-M-1, N-M-2  (5 animals)')
bullet('Soluble Nuclear Naive: N-F-1, N-F-2, N-F-3, N-M-1, N-M-2  (5 animals)')

body('Values are log2-transformed LFQ (Label-Free Quantification) intensities '
     'from MaxQuant, already on the log2 scale in the source file.')

# ══════════════════════════════════════════════════════════════════════════════
# 3. STATISTICAL METHOD
# ══════════════════════════════════════════════════════════════════════════════
heading1('3.  Statistical Method: Delta-Delta Interaction Test')

heading2('3.1  The Core Idea')
body(
    'For each protein, we compute how much it deviated from the naive mean in each '
    'AW animal — separately for the Chromatin and Soluble Nuclear fractions. These '
    'deviations are called delta values. We then test whether the Chromatin deltas '
    'are significantly larger than the Soluble Nuclear deltas across the 4 AW animals.'
)

heading2('3.2  Step-by-Step Calculation')

heading3('Step 1 — Compute the naive baseline for each fraction')
body('For each protein, calculate the mean of all naive replicates in each fraction:')
bullet('mean_CH_naive = mean(Chrom_N-F-1, Chrom_N-F-2, Chrom_N-F-3, Chrom_N-M-1, Chrom_N-M-2)', level=0)
bullet('mean_SN_naive = mean(Nuc_N-F-1, Nuc_N-F-2, Nuc_N-F-3, Nuc_N-M-1, Nuc_N-M-2)', level=0)

heading3('Step 2 — Compute delta values for each AW animal')
body('For each of the 4 AW animals (F-1, F-2, M-1, M-2), subtract the naive mean:')
bullet('delta_CH[i]  =  Chrom_AW[i]  -  mean_CH_naive', level=0)
bullet('delta_SN[i]  =  Nuc_AW[i]    -  mean_SN_naive', level=0)
body(
    'This gives 4 delta_CH values and 4 delta_SN values per protein. Each delta '
    'represents how much that individual AW animal deviated from the naive average '
    'in that fraction. A positive delta means the protein was higher in AW than '
    'in naive; a negative delta means it was lower.',
    indent=False
)

heading3('Step 3 — Test whether Chromatin deltas exceed Soluble Nuclear deltas')
body(
    "Welch's independent-samples t-test is applied comparing the 4 delta_CH values "
    'against the 4 delta_SN values. Welch\'s t-test (rather than Student\'s t-test) '
    'is used because it does not assume equal variance between the two groups, '
    'which is appropriate here since the two fractions may have different variability.'
)
body(
    'The null hypothesis is that the mean deviation from naive is the same in both '
    'fractions. A significant result (after FDR correction) means the protein changed '
    'significantly more in one fraction than the other during AW.'
)

heading3('Step 4 — Direction: Into Chromatin vs Into SN')
body('The interaction score is defined as:')
body('    Interaction score  =  mean(delta_CH)  -  mean(delta_SN)', indent=True)
body(
    'A positive interaction score means the protein increased MORE in Chromatin '
    'than in SN during AW — consistent with translocation INTO Chromatin. '
    'A negative score means the opposite (protein moved into SN). '
    'Only proteins with a positive score and a significant p-value are classified '
    'as "moved into Chromatin."'
)

heading2('3.3  Multiple Testing Correction')
body(
    'Because 2,190 proteins are tested simultaneously, we apply Benjamini-Hochberg '
    '(BH) False Discovery Rate (FDR) correction to the raw p-values. This method '
    'controls the expected proportion of false positives among all proteins called '
    'significant, and is the standard correction for proteomics and genomics experiments.'
)
body(
    'The adjusted p-value threshold used is p_adj < 0.05 (5% FDR). This means that '
    'among all proteins called significant, we expect at most 5% to be false positives.'
)

heading2('3.4  Results Summary')
bullet('Total proteins tested: 2,190  (present in both fractions with sufficient data)')
bullet('Moved significantly INTO Chromatin (p_adj < 0.05, positive interaction score): 346 proteins')
bullet('Moved significantly INTO Soluble Nuclear (p_adj < 0.05, negative interaction score): 362 proteins')

heading2('3.5  Worked Example: Lypla1')
body(
    'To make this concrete, here is the calculation for Lypla1:'
)
bullet('Chromatin naive mean: -4.40  (log2 LFQ)')
bullet('SN naive mean: -4.17  (log2 LFQ)')
bullet('Chromatin AW mean: +4.48  →  delta_CH = 4.48 - (-4.40) = +8.88')
bullet('SN AW mean: -4.21  →  delta_SN = -4.21 - (-4.17) = -0.04')
bullet('Interaction score = 8.88 - (-0.04) = +8.92  (strongly into Chromatin)')
bullet('Welch t-test on [8.88, ...] vs [-0.04, ...]: p_adj = 0.0006')
body(
    'Interpretation: During AW, Lypla1 increased dramatically in the Chromatin '
    'fraction (+8.88 above naive baseline) while remaining essentially flat in the '
    'Soluble Nuclear fraction (-0.04). This is a textbook translocation signal — '
    'the protein physically relocated into the chromatin-bound compartment during '
    'withdrawal and did not simply become more abundant overall.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. SUBGROUP CLASSIFICATION
# ══════════════════════════════════════════════════════════════════════════════
heading1('4.  Subgroup Classification of the 346 Proteins')

body(
    'Among the 346 proteins that moved significantly into Chromatin during AW, '
    'we further classify them by what happened in the Soluble Nuclear fraction '
    'during AW, using the mean log2 FC vs naive (SN_FC_AW):'
)
bullet('Group 1 (G1) — True Translocation  (n = 60): '
       'SN_FC_AW < -0.5.  Chromatin UP and SN DOWN.  '
       'These proteins increased in Chromatin and simultaneously decreased in SN, '
       'the strongest evidence for physical translocation from SN into Chromatin.')
bullet('Group 2 (G2) — Chromatin-selective increase  (n = 206): '
       'SN_FC_AW between -0.5 and +0.5.  Chromatin UP and SN FLAT.  '
       'These proteins increased in Chromatin while SN remained roughly unchanged. '
       'Consistent with selective chromatin binding, but SN pool was not depleted.')
bullet('Group 3 (G3) — General nuclear increase  (n = 80): '
       'SN_FC_AW > +0.5.  Both Chromatin and SN UP.  '
       'These proteins increased in both compartments, suggesting a general increase '
       'in nuclear abundance rather than selective translocation.')

# ══════════════════════════════════════════════════════════════════════════════
# 5. LOG2 FOLD CHANGE VALUES
# ══════════════════════════════════════════════════════════════════════════════
heading1('5.  Log2 Fold Change Values Used in the Heatmaps')

body(
    'The heatmaps display log2 Fold Change (FC) vs naive — not the delta values '
    'used for the statistical test. The FC values are calculated as:'
)
body('    FC  =  mean(condition replicates)  -  mean(naive replicates)  [in log2 space]',
     indent=True)
body(
    'Because the data are already on the log2 scale, this subtraction is equivalent '
    'to log2(condition / naive). A value of +1 means the protein is 2-fold higher '
    'than naive; a value of -1 means half as abundant; 0 means no change.'
)
body('FC values are computed for all three conditions shown in the heatmaps:')
bullet('Intoxication (Intox): mean of I-F-1, I-F-2, I-F-3, I-M-1, I-M-2 vs naive')
bullet('Acute Withdrawal (AW): mean of AW-F-1, AW-F-2, AW-M-1, AW-M-2 vs naive')
bullet('Protracted Abstinence (PA): mean of PA-F-1, PA-F-2, PA-F-3, PA-M-1, PA-M-2 vs naive')

# ══════════════════════════════════════════════════════════════════════════════
# 6. HEATMAP INTERPRETATION
# ══════════════════════════════════════════════════════════════════════════════
heading1('6.  How to Interpret Each Heatmap')

heading2('6.1  General Layout')
body(
    'All heatmaps share the same layout. Each column is one protein. Proteins are '
    'ordered left to right by their Chromatin AW fold change, from highest to lowest. '
    'The two stacked panels — Chromatin (top, cream background) and Soluble Nuclear '
    '(bottom, light blue background) — show the same set of proteins in the same order, '
    'so you can directly compare what happened in each compartment for any given protein '
    '(column) across conditions (rows).'
)
body('Color scale:')
bullet('Deep red / pink  =  protein is HIGHER than naive (positive log2 FC)')
bullet('White  =  no change vs naive (log2 FC near 0)')
bullet('Blue  =  protein is LOWER than naive (negative log2 FC)')
body('The color scale runs from -3 to +3 log2 FC (8-fold range in either direction).')

heading2('6.2  Heatmap A — 346 Proteins, All 3 Conditions')
body(
    'File: Translocation_v3_heatmap_346_pub.pdf'
)
body(
    'This shows all 346 proteins that were statistically identified as moving '
    'significantly more into Chromatin than into SN during AW (p_adj < 0.05, '
    'positive interaction score).'
)
body('What to look for:')
bullet(
    'Chromatin panel, AW row (middle row): predominantly deep red across all proteins. '
    'This is expected and confirms the selection criterion — every protein here increased '
    'significantly more in Chromatin than in SN during AW.'
)
bullet(
    'Chromatin panel, Intox and PA rows: variable signal. Some proteins were already '
    'higher during intoxication; others normalize during protracted abstinence. '
    'This temporal pattern shows the dynamics of the chromatin-bound pool across '
    'the alcohol exposure timeline.'
)
bullet(
    'Soluble Nuclear panel: mixed colors across all rows. Because the 346 proteins '
    'include G1 (SN down), G2 (SN flat), and G3 (SN also up), the SN panel does '
    'not have a consistent direction. This heterogeneity is expected — the only '
    'thing these 346 proteins share is that their Chromatin increase was significantly '
    'larger than their SN change.'
)
bullet(
    'Proteins on the far left: highest Chromatin AW fold change — the most strongly '
    'recruited to chromatin. Proteins on the far right: smallest (but still significant) '
    'Chromatin AW fold change.'
)

heading2('6.3  Heatmap B — 60 True Translocation Proteins, All 3 Conditions')
body(
    'File: Translocation_v3_heatmap_60_pub.pdf'
)
body(
    'This is the most biologically compelling subset: Group 1 proteins (n = 60) that '
    'were significantly higher in Chromatin AND significantly lower in SN during AW '
    '(SN log2 FC < -0.5). These show the clearest evidence of physical movement from '
    'the soluble nuclear pool into the chromatin-bound pool.'
)
body('What to look for:')
bullet(
    'Chromatin panel, AW row: solidly red — these proteins are substantially '
    'elevated in the chromatin fraction during withdrawal.'
)
bullet(
    'Soluble Nuclear panel, AW row: solidly blue — the same proteins are simultaneously '
    'depleted from the soluble nuclear fraction during withdrawal.'
)
bullet(
    'This mirror pattern (Chromatin red + SN blue for the same protein column) is the '
    'hallmark of translocation. The protein pool that was soluble in the nucleus '
    'has shifted into chromatin-bound form during acute withdrawal.'
)
bullet(
    'Intox and PA rows in both panels: generally more variable and less extreme. '
    'The translocation phenomenon is largely specific to the AW timepoint — it does '
    'not persist strongly into protracted abstinence for most proteins, suggesting '
    'the chromatin reorganization is a transient withdrawal response rather than a '
    'permanent epigenetic change.'
)

heading2('6.4  Heatmap C — 346 Proteins, Acute Withdrawal Only')
body(
    'File: Translocation_v3_heatmap_346_AW_pub.pdf'
)
body(
    'A simplified version of Heatmap A showing only the AW timepoint. Each compartment '
    'is reduced to a single row.'
)
body('What to look for:')
bullet(
    'Chromatin row: a smooth gradient from deep red (left) to pale/white (right), '
    'reflecting the sort order (highest to lowest Chromatin AW FC). The entire bar '
    'is in the red-to-white range because all 346 proteins were selected for having '
    'a positive Chromatin fold change during AW.'
)
bullet(
    'Soluble Nuclear row: heterogeneous — blues, whites, and some reds scattered '
    'throughout. This is the mixed G1/G2/G3 composition of the 346-protein group. '
    'The absence of a consistent blue pattern confirms that most of these proteins '
    'did not dramatically deplete from SN — they were recruited to chromatin by '
    'other mechanisms or were newly synthesized.'
)
bullet(
    'Use this heatmap to quickly visualize the scale and heterogeneity of the '
    'chromatin recruitment phenomenon without the complexity of multiple timepoints.'
)

heading2('6.5  Heatmap D — 60 True Translocation Proteins, Acute Withdrawal Only')
body(
    'File: Translocation_v3_heatmap_60_AW_pub.pdf'
)
body(
    'The simplest and most visually striking heatmap: just the AW timepoint '
    'for the 60 true translocation proteins.'
)
body('What to look for:')
bullet(
    'Chromatin row: red gradient from dark red to lighter pink. Every protein '
    'increased in chromatin during withdrawal.'
)
bullet(
    'Soluble Nuclear row: predominantly blue. Every protein in this panel was '
    'selected to have SN_FC_AW < -0.5, so the SN pool was depleted by at least '
    '1.4-fold for all of them.'
)
bullet(
    'The contrast between the two rows — red on top, blue on bottom — is the '
    'visual signature of nuclear-to-chromatin translocation. This figure is '
    'well-suited for a main figure panel or graphical abstract.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. OUTPUT FILES
# ══════════════════════════════════════════════════════════════════════════════
heading1('7.  Output Files')

body('All files are located in:')
body('    .../rat alcohol v2/Translocation SN_Chroamtin/Translocation_v3/', indent=True)

heading2('Statistics')
bullet('Translocation_v3_stats.xlsx — main results file with the following sheets:')
bullet('All_proteins: all 2,190 proteins tested, with raw values, delta values, '
       'interaction scores, p-values, adjusted p-values, and direction', level=1)
bullet('Into_Chromatin: the 346 proteins that moved into Chromatin (p_adj < 0.05), '
       'with Subgroup column (G1/G2/G3)', level=1)
bullet('G1_CH_up_SN_down_n60: 60 true translocation proteins', level=1)
bullet('G2_CH_up_SN_flat_n206: 206 chromatin-selective proteins', level=1)
bullet('G3_CH_up_SN_up_n80: 80 proteins with general nuclear increase', level=1)
bullet('Into_SN: 362 proteins that moved into SN (p_adj < 0.05)', level=1)

heading2('Heatmaps')
bullet('Translocation_v3_heatmap_346_pub.pdf/.png — 346 proteins, 3 conditions')
bullet('Translocation_v3_heatmap_60_pub.pdf/.png — 60 true translocation, 3 conditions')
bullet('Translocation_v3_heatmap_346_AW_pub.pdf/.png — 346 proteins, AW only')
bullet('Translocation_v3_heatmap_60_AW_pub.pdf/.png — 60 true translocation, AW only')

heading2('Scripts')
bullet('make_translocation_v3_stats.py — runs the full statistical analysis and '
       'produces Translocation_v3_stats.xlsx')
bullet('make_translocation_v3_heatmaps_pub.py — generates all four publication heatmaps')

# ══════════════════════════════════════════════════════════════════════════════
# 8. KEY PARAMETERS
# ══════════════════════════════════════════════════════════════════════════════
heading1('8.  Key Parameters at a Glance')

data = [
    ('Statistical test',          "Welch's independent-samples t-test (unequal variance)"),
    ('Comparison',                'delta_CH vs delta_SN (4 values each, per AW animal)'),
    ('Multiple testing correction','Benjamini-Hochberg (BH) FDR'),
    ('Significance threshold',    'p_adj < 0.05'),
    ('Direction filter',          'Interaction score > 0 (Chromatin increase > SN increase)'),
    ('True translocation filter', 'SN_FC_AW < -0.5 (SN depleted by > 1.4-fold)'),
    ('AW animals excluded',       'AW-M-3 (absent from source file)'),
    ('Chromatin filter',          'Keep + Review only (Exclude removed)'),
    ('SN filter',                 'None — all detected proteins included'),
    ('Heatmap color scale',       'Log2 FC vs naive, range -3 to +3'),
    ('Heatmap sort order',        'Chromatin AW FC, high to low (left to right)'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Parameter'
hdr[1].text = 'Value'
for cell in hdr:
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

for param, value in data:
    row = table.add_row().cells
    row[0].text = param
    row[1].text = value
    for cell in row:
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# 9. HOW TO DESCRIBE THIS IN A PAPER
# ══════════════════════════════════════════════════════════════════════════════
heading1('9.  Suggested Methods Language for a Paper')

body(
    'The following text can be adapted for the Methods section of a manuscript:'
)
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.4)
p.paragraph_format.right_indent = Inches(0.4)
p.paragraph_format.space_after  = Pt(8)
run = p.add_run(
    'To identify proteins that translocated into the chromatin-bound nuclear fraction '
    'during acute alcohol withdrawal, we applied a delta-delta interaction test. For '
    'each protein detected in both the soluble nuclear (SN) and chromatin (CH) fractions, '
    'we computed per-animal deviations from the naive mean (delta_CH and delta_SN) using '
    'the four available AW replicates (AW-F-1, AW-F-2, AW-M-1, AW-M-2; AW-M-3 excluded). '
    "Welch's independent-samples t-test was applied comparing delta_CH versus delta_SN "
    'across the AW animals. P-values were corrected for multiple comparisons using the '
    'Benjamini-Hochberg false discovery rate procedure. Proteins with p_adj < 0.05 and a '
    'positive interaction score (mean delta_CH > mean delta_SN) were classified as moved '
    'into chromatin during AW (n = 346). A subset of 60 proteins that additionally showed '
    'a decrease in SN (SN log2 FC vs naive < -0.5) were classified as true translocation '
    'proteins, indicating concurrent depletion of the soluble nuclear pool and accumulation '
    'in the chromatin-bound fraction.'
)
run.italic = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(50, 50, 50)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_DOCX)
print(f'Saved: {OUT_DOCX}')
