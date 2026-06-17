"""
Generate Translocation_Analysis_Summary.docx using python-docx.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/naomi/Desktop/rat alcohol v2/Translocation SN_Chroamtin"
OUT  = os.path.join(BASE, "Translocation_Analysis_Summary.docx")

# Image info: filename, (width_px, height_px)
IMAGES = {
    "barplot.jpg":           (496, 349),
    "scatter_fig2_page1.jpg": (800, 750),
    "transloc_p2.jpg":       (567, 494),
    "transloc_p3.jpg":       (785, 351),
}

MAX_WIDTH_INCHES = 6.0

def scaled_dims(w_px, h_px, max_w=MAX_WIDTH_INCHES):
    """Return (width_in, height_in) scaled to fit within max_w."""
    aspect = h_px / w_px
    w = min(max_w, w_px / 96)  # treat 96 dpi as baseline
    # Always cap to max_w
    w = max_w if w_px / 96 > max_w else w_px / 96
    h = w * aspect
    return w, h

def add_run_with_font(para, text, size_pt, bold=False, italic=False):
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    return run

def add_heading(doc, text, size_pt, centered=False):
    para = doc.add_paragraph()
    if centered:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_font(para, text, size_pt, bold=True)
    return para

def add_caption(doc, text):
    para = doc.add_paragraph()
    add_run_with_font(para, text, 10, italic=True)
    return para

def add_body(doc, text):
    para = doc.add_paragraph()
    add_run_with_font(para, text, 11)
    return para

def add_fig_heading(doc, text):
    para = doc.add_paragraph()
    add_run_with_font(para, text, 14, bold=True)
    return para

def add_image(doc, fname):
    w_px, h_px = IMAGES[fname]
    w_in, h_in = scaled_dims(w_px, h_px)
    path = os.path.join(BASE, fname)
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(path, width=Inches(w_in), height=Inches(h_in))
    return para

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)
    return para

doc = Document()

# Set default font for the document body style
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# ── TITLE ──────────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_with_font(title_para,
    "Chromatin vs Soluble Nuclear Compartment Analysis \u2014 Acute Withdrawal",
    16, bold=True)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_with_font(subtitle_para, "Rat Amygdala Proteomics", 12)

doc.add_paragraph()  # spacer

# ── FIGURE 1 ───────────────────────────────────────────────────────────────────
add_fig_heading(doc, "Figure 1: Average Chromatin/Nuclear Protein Ratio by Condition")
add_image(doc, "barplot.jpg")
add_caption(doc,
    "Figure 1. Average fold change (Chromatin \u2212 Nuclear, log\u2082) across 1,115 "
    "AW-significant proteins per condition \u00b1 SEM. One-way ANOVA p = 0.39.")
add_body(doc,
    "What this shows: Each bar is the average difference between a protein\u2019s abundance "
    "in Chromatin versus Soluble Nuclear (in log\u2082 space) for each condition. All bars "
    "are negative, meaning proteins are on average more abundant in the Soluble Nuclear "
    "fraction than in Chromatin across all conditions. During Acute Withdrawal (AW), the "
    "bar is closest to zero (\u22120.07), meaning the gap between the two fractions narrows "
    "the most during AW compared to Naive (\u22120.20) or Protracted Abstinence (\u22120.25). "
    "However, the one-way ANOVA p = 0.39 is not significant \u2014 this global average "
    "dilutes the signal because individual proteins shift in different directions.")

# ── PAGE BREAK ─────────────────────────────────────────────────────────────────
add_page_break(doc)

# ── FIGURE 2 ───────────────────────────────────────────────────────────────────
add_fig_heading(doc, "Figure 2: Scatter Plot \u2014 \u0394Chromatin vs \u0394Nuclear During AW")
add_image(doc, "scatter_fig2_page1.jpg")
add_caption(doc,
    "Figure 2. Each dot is one protein (n = 1,115 AW-significant proteins). "
    "X-axis: \u0394Soluble Nuclear (AW \u2212 Na\u00efve LFQ). "
    "Y-axis: \u0394Chromatin (AW \u2212 Na\u00efve LFQ). "
    "Dashed diagonal = y = x (equal change in both compartments). "
    "Pink dots (n = 429): proteins that changed significantly MORE in Chromatin than Nuclear "
    "(above diagonal, BH-corrected p < 0.05). "
    "Blue dots (n = 672): proteins that changed significantly MORE in Nuclear than Chromatin "
    "(below diagonal, BH-corrected p < 0.05). "
    "Pearson r = 0.440 (p = 4.1\u00d710\u207b\u2075\u2074); Spearman \u03c1 = 0.460 (p = 1.9\u00d710\u207b\u2075\u2079). "
    "Binomial test on directional split (429 vs 672): p = 2.4\u00d710\u207b\u00b9\u00b3.")
add_body(doc,
    "What this shows: The two axes represent how much each protein\u2019s abundance changed "
    "during Acute Withdrawal (AW) relative to Naive \u2014 separately in the Chromatin fraction "
    "(y-axis) and the Soluble Nuclear fraction (x-axis). Each protein is placed at its "
    "(\u0394Nuclear, \u0394Chromatin) coordinate. "
    "\n\nThe dashed diagonal (y = x) is the line of equal change: a protein sitting exactly "
    "on it changed by the same amount in both compartments. Points above the diagonal "
    "changed MORE in Chromatin; points below changed MORE in Nuclear. "
    "\n\nThree key statistics support this figure:"
    "\n\n1. Pearson r = 0.440 (p = 4.1\u00d710\u207b\u2075\u2074): The two compartments are "
    "significantly positively correlated \u2014 when a protein goes up in Chromatin it tends "
    "to also go up in Nuclear. This is the p-value for the overall scatter plot relationship."
    "\n\n2. Per-protein BH-corrected Welch\u2019s t-test: For each protein, we tested whether "
    "its change in Chromatin was significantly different from its change in Nuclear across "
    "the AW replicates. 429 proteins were significantly enriched in Chromatin (pink, above "
    "diagonal) and 672 were significantly enriched in Nuclear (blue, below diagonal), both "
    "at BH-corrected p < 0.05."
    "\n\n3. Binomial test (p = 2.4\u00d710\u207b\u00b9\u00b3): Of the 1,101 proteins with a "
    "significant compartment preference, 672 shifted toward Nuclear and 429 toward Chromatin. "
    "A binomial test confirms this 61:39 split is significantly different from a random 50:50 "
    "distribution, meaning proteins during AW preferentially accumulate in the Soluble Nuclear "
    "fraction rather than Chromatin."
    "\n\nImportant: \u2018INTO Chromatin\u2019 means the protein\u2019s change was significantly "
    "larger in Chromatin than in Soluble Nuclear during AW \u2014 it does not necessarily mean "
    "the protein disappeared from Nuclear. The 429 Chromatin-enriched proteins are candidates "
    "for AW-specific chromatin association and are suitable for pathway analysis (e.g., Metascape). "
    "Only 11 proteins showed strict bidirectional redistribution (UP in Chromatin AND DOWN in "
    "Nuclear simultaneously).")

doc.add_paragraph()  # spacer

# ── FIGURE 3 ───────────────────────────────────────────────────────────────────
add_fig_heading(doc, "Figure 3: Translocation Volcano Plot \u2014 Acute Withdrawal")
add_image(doc, "transloc_p2.jpg")
add_caption(doc,
    "Figure 3. Volcano plot of the interaction score vs statistical significance. "
    "Pink = 429 proteins significantly enriched in Chromatin during AW (BH p < 0.05). "
    "Blue = 672 proteins significantly enriched in Nuclear. Dashed line = p = 0.05.")
add_body(doc,
    "What this shows: The x-axis is the interaction score (\u0394Chromatin \u2212 "
    "\u0394Nuclear) \u2014 how much more a protein changed in Chromatin versus Nuclear "
    "during AW. The y-axis is statistical significance. Pink dots (right side) are proteins "
    "where Chromatin showed a significantly greater increase than Nuclear during AW. Blue "
    "dots (left) showed the opposite. More proteins shifted toward Nuclear (672) than "
    "Chromatin (429), but the 429 Chromatin-enriched proteins are biologically interesting "
    "as candidates for AW-specific chromatin association. Important: \u2018shifted toward "
    "Chromatin\u2019 means the Chromatin increase was greater than the Nuclear increase "
    "relative to Naive \u2014 it does not necessarily mean Nuclear decreased.")

doc.add_paragraph()  # spacer

# ── FIGURE 4 ───────────────────────────────────────────────────────────────────
add_fig_heading(doc, "Figure 4: Cluster Tests \u2014 Non-Circular Translocation Analysis")
add_image(doc, "transloc_p3.jpg")
add_caption(doc,
    "Figure 4. Left (pink): Distribution of \u0394Nuclear values for the 362 proteins "
    "significantly upregulated in Chromatin during AW (mean = +0.49, p = 8.83\u00d710\u207b\u00b9\u00b3). "
    "Right (blue): Distribution of \u0394Chromatin values for the 150 proteins significantly "
    "downregulated in Nuclear during AW (mean = \u22121.00, p = 1.42\u00d710\u207b\u2077).")
add_body(doc,
    "What this shows: These histograms use a non-circular approach to test the translocation "
    "hypothesis. Left panel (Cluster A): We took the 362 proteins significantly UP in "
    "Chromatin during AW and asked whether their Nuclear levels also changed. The mean "
    "\u0394Nuclear = +0.49 (right of zero), meaning these proteins also INCREASED in Nuclear "
    "\u2014 confirming global upregulation rather than translocation from Nuclear to Chromatin. "
    "Right panel (Cluster B): We took the 150 proteins significantly DOWN in Nuclear and "
    "asked whether their Chromatin levels went up. The mean \u0394Chromatin = \u22121.00 "
    "(left of zero) \u2014 their Chromatin also decreased. Again consistent with global "
    "downregulation. Conclusion: The dominant signal in this dataset is global protein "
    "abundance changes rather than redistribution between compartments. The 429 proteins "
    "showing greater Chromatin enrichment during AW likely reflect AW-specific chromatin "
    "association, which is still biologically meaningful. Only 11 proteins showed strict "
    "bidirectional redistribution (up in Chromatin AND down in Nuclear simultaneously).")

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"File size: {os.path.getsize(OUT):,} bytes")
