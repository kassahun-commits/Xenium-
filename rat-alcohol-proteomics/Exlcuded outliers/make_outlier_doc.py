from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = 'AW_M3_Exclusion_Methods.docx'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def heading(text, size=14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'; r.font.size = Pt(size); r.font.bold = True
    return p

def body(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'; r.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Arial'; r.font.size = Pt(11)
    return p

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Exclusion of AW-M-3 Replicate — Rationale and Methods')
r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True
doc.add_paragraph()

heading('Summary')
body(
    'During quality control of the Acute Withdrawal (AW) proteomics data, '
    'the replicate sample AW-M-3 was identified as a systematic outlier across '
    'all four subcellular compartments (Membrane, Cytosol, Chromatin, Soluble Nuclear) '
    'and was excluded from all downstream analyses. All other AW replicates '
    '(AW-F-1, AW-F-2, AW-M-1, AW-M-2) were retained.'
)
doc.add_paragraph()

heading('What We Found')
body(
    'The median LFQ intensity of each AW replicate was calculated across all proteins '
    'in each compartment. In every compartment, AW-M-3 had a median 2–3x lower than '
    'all other AW replicates:'
)
doc.add_paragraph()

# Table-like bullet list
for comp, m3, others in [
    ('Membrane',        '−4.48', '−0.81 to −1.04'),
    ('Cytosol',         '−2.82', '−0.66 to −0.82'),
    ('Chromatin',       '−5.76', '−1.90 to −2.42'),
    ('Soluble Nuclear', '−4.73', '−1.11 to −1.71'),
]:
    bullet(f'{comp}:  AW-M-3 median = {m3}   vs   other AW replicates = {others}')

doc.add_paragraph()
body(
    'This pattern — where one sample is uniformly lower than all others across every '
    'compartment simultaneously — is the hallmark of a failed or degraded sample run, '
    'not a biological effect. A true biological outlier would affect some proteins '
    'differently from others; a global downward shift of all proteins equally indicates '
    'a technical issue such as sample degradation, loading error, or instrument '
    'performance failure during that specific run.'
)
doc.add_paragraph()

heading('Why the Stripe Artifact Persisted After Exclusion')
body(
    'After excluding AW-M-3, a visual "stripe" (vertical band of proteins at similar '
    'fold change values, around FC = −1.3 to −1.5) remained in the Chromatin AW volcano '
    'plot. Investigation showed this is NOT a technical artifact — it reflects the true '
    'biological distribution of the data:'
)
bullet(
    'The FC distribution for Chromatin AW has a genuine peak at approximately −1.3 to −1.5, '
    'with ~276 proteins (out of 2,192) in that narrow range — roughly 3x more than '
    'surrounding FC bins.'
)
bullet(
    'Per-sample median normalization did not remove the peak; it only shifted it '
    'slightly (from −1.3 to −1.5), confirming the signal is intrinsic to the data.'
)
bullet(
    'The same peak is not present in Membrane or Cytosol, which show smoothly '
    'distributed FC values centered near 0.'
)
body(
    'Interpretation: Many chromatin-associated proteins decrease by approximately '
    'the same magnitude during Acute Withdrawal relative to Naïve. This likely '
    'reflects a coordinated, condition-specific reorganisation of the chromatin-bound '
    'proteome during AW — a biologically meaningful finding rather than a technical artifact. '
    'A density volcano plot (VP_Chromatin_AW_density.pdf) is provided alongside the '
    'standard volcano to honestly visualise this distribution, where darker hexbins '
    'indicate more proteins at that fold change / p-value combination.'
)
doc.add_paragraph()

heading('How the Stripe Was Handled in the Panel Figure')
body(
    'Because the stripe represents a real biological signal (not a technical error), '
    'it cannot and should not be removed from the data. However, when hundreds of dots '
    'stack directly on top of each other at the same x-position, the plot can look '
    'artificially banded and hard to read. To improve visual clarity without altering '
    'any statistical results, a small random horizontal jitter was applied to each dot '
    'for display purposes only:'
)
bullet(
    'Each protein\'s x-axis position (Log2 Fold Change) was shifted by a random value '
    'drawn uniformly from ±0.20 or ±0.40 (two versions were produced for comparison). '
    'This spreads the overlapping dots apart so the density of the cloud is visible '
    'rather than hidden behind a single line.'
)
bullet(
    'The jitter is purely cosmetic — it is applied only to the plotted dot positions. '
    'All significance thresholds (corrected p > 3.3, |FC| > 0.5) and protein counts '
    'in the legend are calculated from the original, un-jittered fold change values.'
)
bullet(
    'The same jitter seed (numpy random seed = 42) and the same jitter magnitude were '
    'applied identically to all 12 panels, so no individual compartment or condition '
    'is treated differently from any other.'
)
body(
    'In short: the stripe is still there biologically — the jitter just prevents hundreds '
    'of dots from stacking into a single vertical line that looks like an artifact. '
    'The protein counts in every legend are the ground truth.'
)
doc.add_paragraph()

heading('Impact of AW-M-3 Exclusion on Protein Counts')
body('Excluding AW-M-3 changed the number of significant proteins as follows '
     '(corrected p > 3.3, |FC| > 0.5):')
doc.add_paragraph()
for comp, up_before, dn_before, up_after, dn_after in [
    ('Membrane',        115,  59,  115,  76),
    ('Cytosol',          31,  97,   63,  75),
    ('Chromatin',       362, 656,  433, 642),
    ('Soluble Nuclear', 127, 150,  173, 199),
]:
    bullet(
        f'{comp}:  '
        f'↑ {up_before}→{up_after}  ↓ {dn_before}→{dn_after}'
    )

doc.add_paragraph()
body(
    'The increase in upregulated proteins after AW-M-3 exclusion is expected: '
    'the outlier replicate was globally depressed, artificially pulling the AW '
    'condition mean downward and masking proteins that genuinely increased during AW.'
)
doc.add_paragraph()

heading('Files Generated')
bullet('VP_AW_noAWM3_FixedAxes.pdf — volcano plots for all 4 compartments, AW-M-3 excluded, axes fixed at FC ±7 and corrected p 0–20')
bullet('AW_M3_outlier_check.pdf — diagnostic boxplots and median bar charts showing AW-M-3 as a clear outlier across all compartments')
bullet('VP_Chromatin_AW_density.pdf — hexbin density volcano for Chromatin AW explaining the stripe as a real density peak')
bullet('VP_Panel_Figure_v1_jitter40.pdf — 4×3 panel volcano figure (all compartments × conditions), jitter ±0.40')
bullet('VP_Panel_Figure_v2_jitter20.pdf — 4×3 panel volcano figure (all compartments × conditions), jitter ±0.20')
bullet('AW_M3_Exclusion_Methods.docx — this document')

doc.save(OUT)
print(f'Saved: {OUT}')
