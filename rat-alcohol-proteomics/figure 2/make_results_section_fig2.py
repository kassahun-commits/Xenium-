"""
Figure 2 Results Section — Word Document Generator
====================================================
Creates Results_Section_Figure2.docx in the figure 2 folder.
Sections: Membrane | Cytosol  (more to be added)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DOCX   = os.path.join(SCRIPT_DIR, 'Results_Section_Figures1_2.docx')

doc = Document()

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font  = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Narrow margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name  = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def spacer(doc):
    p = doc.add_paragraph('')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

# ═════════════════════════════════════════════════════════════════════════════
# FIGURE 1 — Experimental design + heatmap overview
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, 'Compartment-Resolved Proteomic Profiling of the Rat Amygdala Across Alcohol Exposure States', level=1)
spacer(doc)

body(doc,
    'To investigate compartment-specific proteomic remodeling induced by chronic ethanol '
    'exposure, we focused on the amygdala, a brain region critically implicated in the '
    'negative-reinforcement processes that drive the progression of AUD [16]. To ensure '
    'consistent and physiologically relevant intoxication, rats underwent a chronic alcohol '
    'vapor chamber paradigm for 7 weeks. Amygdala tissue was collected at three time points: '
    'during intoxication (chronic exposure), 24 h after the final exposure (acute withdrawal), '
    'and 30 d after the last vapor session (protracted abstinence). Amygdala tissue from each '
    'group was harvested for matched proteomic and transcriptomic analyses. For proteomic '
    'profiling, tissue was fractionated into cytosolic, membrane, soluble nuclear, and chromatin '
    'compartments and analyzed by mass spectrometry (MS). Differential expression, network-based '
    'analyses, and ontology enrichment were then used to define compartment-specific and '
    'persistent proteomic signatures associated with each stage of alcohol exposure (Fig. 1A).'
)

body(doc,
    'To provide a global view of proteomic regulation, we identified proteins significantly '
    'changed in at least one condition across each subcellular compartment (union set; corrected '
    'p > 3.3, |log₂ FC| > 0.5 relative to naïve controls) and visualized their fold changes '
    'across all three conditions as a heatmap (Fig. 1B). The number of union-significant proteins '
    'differed markedly across compartments: 383 in the cytosolic fraction, 433 in the membrane '
    'fraction, 570 in the soluble nuclear fraction, and 1,205 in the chromatin fraction. The '
    'disproportionately large number of regulated chromatin-associated proteins suggests that '
    'chronic ethanol exposure induces widespread reorganization of the chromatin-bound proteome, '
    'consistent with broad transcriptional and epigenetic remodeling in the amygdala.'
)

body(doc,
    'Across all compartments, the pattern of regulation was condition-dependent and '
    'compartment-specific, underscoring the value of subcellular fractionation for resolving '
    'alcohol-induced neuroadaptations (Fig. 1B). In the cytosolic fraction, intoxication was '
    'associated with a predominantly upregulated protein signature, while acute withdrawal '
    'produced a striking reversal toward broad protein downregulation, the most pronounced '
    'directional shift observed in this compartment. Protracted abstinence showed an intermediate '
    'pattern with a mixture of increased and decreased proteins. The membrane fraction exhibited '
    'a similar directionality during intoxication, with upregulation dominating, followed by a '
    'progressive increase in downregulated proteins across withdrawal and abstinence, the latter '
    'being characterized by the largest proportion of decreased membrane-associated proteins.'
)

body(doc,
    'The soluble nuclear fraction displayed the most dramatic acute withdrawal response across '
    'all compartments, with a deep and sustained shift toward protein downregulation during '
    'withdrawal that partially recovered at protracted abstinence. This pronounced withdrawal '
    'signature in the soluble nuclear fraction may reflect acute transcriptional reprogramming '
    'or disrupted nuclear import of regulatory proteins during the withdrawal state. In contrast, '
    'the chromatin-associated fraction — the largest and most extensively regulated compartment '
    '— showed a more complex and heterogeneous pattern, with proteins both increasing and '
    'decreasing across all three conditions, consistent with condition-specific remodeling of '
    'chromatin-bound regulatory factors throughout the course of alcohol exposure and abstinence. '
    'Together, these data reveal that chronic alcohol exposure engages distinct and temporally '
    'dynamic proteomic programs within each subcellular compartment of the amygdala, motivating '
    'detailed compartment-specific analyses of the biological processes underlying these changes.'
)

spacer(doc)

# ═════════════════════════════════════════════════════════════════════════════
# FIGURE 2 HEADER
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, 'Compartment-Specific Biological Processes Regulated by Alcohol Exposure', level=1)
spacer(doc)

# ═════════════════════════════════════════════════════════════════════════════
# MEMBRANE
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, 'Membrane Proteome', level=2)

body(doc,
    'Volcano plot analysis of the amygdala membrane-associated proteome revealed '
    'condition-dependent shifts in protein abundance across all three alcohol exposure states '
    '(Fig. 2A). During intoxication, 92 proteins were significantly increased and 80 were '
    'decreased relative to naïve controls. Acute withdrawal produced the most expansive '
    'bidirectional response in the membrane fraction, with 115 proteins increased and 76 '
    'decreased. Protracted abstinence was distinguished by a pronounced shift toward '
    'downregulation, with 137 proteins decreased compared to only 47 increased, suggesting '
    'progressive membrane proteome reorganization over the course of abstinence.'
)

body(doc,
    'Gene Ontology (GO) network analysis of upregulated membrane proteins during intoxication '
    'revealed hemostasis as a significantly enriched biological process (Fig. 2A). In the brain, '
    'many proteins annotated to this pathway contribute to cytoskeletal organization, '
    'microtubule-based transport, G-protein regulation, and membrane-associated signal '
    'transduction, rather than canonical vascular repair [CITE]. Within this module, we observed '
    'increased abundance of Superoxide Dismutase 1 (Sod1) and Phosphodiesterase 1B (Pde1b), two '
    'proteins with prior links to alcohol-related phenotypes. Sod1 is a Cu/Zn superoxide '
    'dismutase that detoxifies superoxide anions and contributes to antioxidant defense; altered '
    'Sod1 expression and activity have been repeatedly reported after chronic alcohol exposure and '
    'in alcohol use disorder [17,18]. Pde1b encodes a Ca2+/calmodulin-regulated phosphodiesterase '
    'that hydrolyzes cAMP and cGMP, thereby shaping cyclic-nucleotide signaling in dopamine-rich '
    'brain regions, and has been implicated in regulating alcohol drinking behaviors in rodent '
    'models [19]. Enrichment of Sod1 and Pde1b in the amygdala membrane fraction therefore aligns '
    'with prior evidence that oxidative stress defenses and cyclic-nucleotide signaling pathways '
    'are engaged in alcohol use disorder-related neuroadaptations.'
)

body(doc,
    'Cocaine addiction also emerged as a significantly enriched term among upregulated membrane '
    'proteins during intoxication. Proteins annotated to this pathway are implicated more broadly '
    'in substance use disorders, as most drugs of abuse converge on disruption of neuronal '
    'plasticity through G-protein signaling cascades and downstream effectors [20]. Among the '
    'upregulated proteins contributing to this module were Glutamate ionotropic receptor NMDA type '
    'subunit 2B (Grin2b), which encodes an NMDA receptor subunit long tied to alcohol\'s effects '
    'on excitatory transmission and dependence phenotypes [21,22,23], and Regulating synaptic '
    'membrane exocytosis protein 1 (Rims1), an active-zone scaffold that recruits and activates '
    'Munc13-1 to facilitate synaptic vesicle priming at the presynaptic membrane. Downregulated '
    'membrane proteins during intoxication were enriched for regulation of translation and '
    'cellular response to stress, suggesting a suppression of membrane-associated adaptive '
    'stress-response machinery during acute alcohol exposure.'
)

body(doc,
    'During acute withdrawal, upregulated membrane proteins were enriched for oxidation of organic '
    'compounds and purine nucleotide metabolic processes, pointing to heightened oxidative '
    'metabolism and energy demand at the membrane during this state. Downregulated proteins were '
    'associated with regulation of actin cytoskeleton organization and vesicle-mediated transport, '
    'indicating a disruption of membrane trafficking and cytoskeletal remodeling coinciding with '
    'the peak of the withdrawal response.'
)

body(doc,
    'Protracted abstinence was characterized by a predominance of protein downregulation in the '
    'membrane fraction. Upregulated proteins were enriched for regulation of exocytosis and '
    'regulation of neutrophil chemotaxis, the latter reflecting immune-related signaling pathways '
    'increasingly implicated in the neuroinflammatory dimension of alcohol dependence [CITE]. '
    'Downregulated proteins were associated with metabolism of lipids and regulation of '
    'protein-containing complex assembly, consistent with broad reorganization of membrane lipid '
    'composition and multiprotein signaling complexes during long-term abstinence.'
)

spacer(doc)

# ═════════════════════════════════════════════════════════════════════════════
# CYTOSOL
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, 'Cytosolic Proteome', level=2)

body(doc,
    'We next interrogated alcohol-induced changes to the cytosolic proteome. GO network analysis '
    'of upregulated cytosolic proteins revealed enrichment of amino acid metabolic processes '
    '(Fig. 2B). Proteins associated with this term include glutamate cysteine ligase modifier '
    'subunit (Gclm), which was significantly increased in the alcohol exposure group. As the '
    'regulatory subunit of the rate-limiting enzyme for glutathione synthesis, Gclm supports redox '
    'buffering in neurons and glia [25]. Prior work in individuals with AUD reports altered brain '
    'glutathione levels, and our findings align with this pattern as Gclm is involved in '
    'glutathione production [26]. Long-chain acyl-CoA synthetase 1 (Acsl1) was another protein '
    'upregulated within the amino acid metabolic process network. Acsl1 activates long-chain fatty '
    'acids to acyl-CoA for entry into lipid metabolic pathways [27]. Whereas chronic alcohol '
    'suppresses hepatic ACSL1 in alcohol-associated liver disease, the increase we detect in the '
    'brain indicates tissue-specific metabolic reprogramming toward enhanced activation and flux of '
    'long-chain fatty acids [28].'
)

body(doc,
    'Dicarboxylate metabolism was another process enriched in the cytosolic fraction, indicating '
    'increased handling of metabolites such as fumarate and malate that feed the TCA cycle and '
    'couple carbon flow to nitrogen disposal [29]. In neural cells, this pathway supports '
    'anaplerosis to sustain ATP production, helps rebalance redox equivalents, and links amino '
    'acid breakdown to urea cycle function [30]. Two enzymes on our list are strongly linked to '
    'this process: Argininosuccinate lyase (ASL), which generates fumarate from '
    'argininosuccinate, and 4-Hydroxyphenylpyruvate dioxygenase (HPD), a tyrosine-catabolic '
    'enzyme that ultimately provides fumarate to the TCA cycle [31,32].'
)

body(doc,
    'Notably, alcohol dehydrogenase 5 (Adh5), a brain-expressed alcohol-metabolizing enzyme, was '
    'significantly increased in the cytosolic fraction. Prior studies have reported '
    'alcohol-induced increases in Cyp2e1 levels and catalase activity in neural and peripheral '
    'tissues; however, to our knowledge, there are no reports of Adh5 upregulation in the brain '
    'following chronic alcohol exposure [33]. Our finding therefore suggests a previously '
    'unrecognized adjustment of the cerebral alcohol metabolism network, in which Adh5 may '
    'contribute to alcohol and aldehyde handling alongside canonical pathways. Adh5 also serves '
    'as a redox enzyme, functioning as the glutathione-dependent formaldehyde dehydrogenase and '
    'S-nitrosoglutathione reductase — processes essential for detoxication of reactive metabolites '
    '[34]. Further analyses are warranted to determine whether Adh5 induction primarily supports '
    'alcohol metabolism in the amygdala or instead reflects a compensatory response to the redox '
    'stress associated with chronic alcohol exposure [35].'
)

spacer(doc)

# ═════════════════════════════════════════════════════════════════════════════
# SOLUBLE NUCLEAR
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, 'Soluble Nuclear Proteome', level=2)

body(doc,
    'Within the soluble nuclear fraction, proteins increased following chronic ethanol exposure '
    'were enriched for the biological processes dicarboxylic acid metabolic process and small '
    'molecule catabolic process, indicating altered nuclear-associated metabolic activity '
    '(Fig. 2C). These pathways broadly include enzymes involved in amino acid, organic acid, and '
    'intermediary metabolite turnover that can influence redox balance, energy state, and '
    'availability of substrates linked to transcriptional regulation.'
)

body(doc,
    'For the dicarboxylic acid metabolic process, two notable proteins were MTHFD1 and IDH2. '
    'MTHFD1 is a one-carbon metabolism enzyme with established nuclear functions that contributes '
    'to folate-mediated transfer of one-carbon units, supporting nucleotide synthesis and methyl '
    'donor pathways relevant to chromatin regulation. IDH2, a key regulator of mitochondrial and '
    'cellular redox metabolism, generates metabolites that can influence epigenetic enzyme '
    'activity and oxidative stress responses. Their increased abundance in the soluble nuclear '
    'compartment suggests chronic ethanol exposure may enhance metabolic pathways that support '
    'nuclear remodeling and transcriptional adaptation.'
)

body(doc,
    'Similarly, enrichment of the small molecule catabolic process was represented by proteins '
    'such as ENO1, a glycolytic enzyme with reported stress-responsive nuclear functions, and '
    'ALDH1A1, which participates in aldehyde detoxification and retinoid metabolism. Together, '
    'these findings support a model in which chronic ethanol exposure reorganizes soluble nuclear '
    'metabolism, potentially linking cellular metabolic state to gene regulatory responses in '
    'the amygdala.'
)

body(doc,
    'Downregulated soluble nuclear proteins were enriched for [PLACEHOLDER — GO terms to be '
    'added] and [PLACEHOLDER — GO term 2]. [PLACEHOLDER — brief interpretive sentence '
    'connecting the downregulated processes to nuclear function and alcohol exposure.]'
)

spacer(doc)

# ═════════════════════════════════════════════════════════════════════════════
# CHROMATIN
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, 'Chromatin-Associated Proteome', level=2)

body(doc,
    'Proteomic analysis of the chromatin-enriched fraction from rat amygdala following chronic '
    'ethanol exposure identified small molecule catabolic process as the top enriched biological '
    'pathway (Fig. 2D). This term encompasses enzymes involved in the breakdown and '
    'interconversion of low-molecular-weight metabolites that regulate cellular energy state, '
    'redox balance, and substrate availability. Their enrichment within the chromatin compartment '
    'suggests that chronic ethanol exposure may promote recruitment of metabolic proteins to the '
    'nucleus, where they could influence transcriptional activity and epigenetic regulation '
    'through local metabolite production.'
)

body(doc,
    'Several proteins within this category have prior relevance to nuclear function. GAPDH and '
    'ENO1, classically glycolytic enzymes, have been reported to translocate to the nucleus under '
    'cellular stress and participate in transcriptional regulation. AHCY, a key regulator of the '
    'methionine cycle, controls cellular methylation potential and may impact histone or DNA '
    'methylation states. MTHFD1, a one-carbon metabolism enzyme with known nuclear roles, '
    'supports pathways linked to nucleotide synthesis and methyl donor availability. Together, '
    'these findings suggest chronic ethanol exposure may remodel chromatin in part through '
    'recruitment of metabolic enzymes that couple cellular metabolism to gene regulation.'
)

body(doc,
    'The second enriched pathway, valine, leucine, and isoleucine degradation, reflects '
    'branched-chain amino acid catabolism. Enrichment of this pathway in the chromatin fraction '
    'may indicate altered amino acid metabolic signaling within nuclear compartments after '
    'ethanol exposure. Representative proteins included ALDH2, involved in aldehyde '
    'detoxification and oxidative stress responses, and HMGCS2, which links leucine catabolism '
    'to ketone and acetyl-CoA metabolism. These data support a broader model in which chronic '
    'ethanol exposure reorganizes metabolic networks at chromatin to influence transcriptional '
    'adaptation in the amygdala.'
)

body(doc,
    'In contrast to the proteins increased in the chromatin fraction, chronic ethanol '
    'intoxication also produced a set of downregulated chromatin-associated proteins enriched '
    'for the biological processes regulation of exocytosis and neuronal signaling. These terms '
    'generally reflect proteins involved in calcium-dependent vesicle release, membrane '
    'trafficking, and signal transduction pathways that control neurotransmitter secretion and '
    'synaptic communication. Their reduced abundance within the chromatin fraction may indicate '
    'diminished recruitment of activity-responsive signaling proteins to nuclear compartments '
    'during intoxication.'
)

body(doc,
    'Several proteins within these pathways have known signaling functions that can extend to '
    'the nucleus. PRKCG and PRKCB are protein kinase C isoforms that participate in '
    'stimulus-dependent signaling cascades and have been implicated in nuclear signaling '
    'responses. CDK5, a major neuronal kinase, has documented nuclear localization and regulates '
    'activity-dependent gene expression. In addition, RAB11B and STXBP3 are associated with '
    'vesicle trafficking and membrane recycling, suggesting reduced coupling between synaptic '
    'vesicle dynamics and transcriptional programs. Collectively, these findings are consistent '
    'with suppression of neuronal activity-linked chromatin signaling pathways during chronic '
    'ethanol intoxication.'
)

spacer(doc)

# ═════════════════════════════════════════════════════════════════════════════
# FIGURE 3 — Chromatin AW: volcano plots, Venn diagrams, GO terms
# ═════════════════════════════════════════════════════════════════════════════
heading(doc, 'Acute Withdrawal Drives Extensive and Largely State-Specific Remodeling of the Chromatin-Associated Proteome', level=1)
spacer(doc)

body(doc,
    'The union heatmap revealed that the chromatin compartment harbored the largest and most '
    'dynamic proteomic response across all three alcohol exposure states (Fig. 1B). The acute '
    'withdrawal condition was particularly conspicuous, displaying both the broadest upregulation '
    'and the most extensive downregulation of chromatin-associated proteins relative to naïve '
    'controls. To characterize this withdrawal-specific remodeling in detail, we generated '
    'volcano plots separately highlighting proteins increased and decreased during acute '
    'withdrawal (Fig. 3A–B). These plots revealed that 433 chromatin-associated proteins were '
    'significantly increased and 642 were significantly decreased during acute withdrawal, out of '
    '2,192 proteins quantified in this compartment, representing one of the largest bidirectional '
    'regulatory responses observed across all conditions and compartments in this dataset.'
)

body(doc,
    'To determine the extent to which the withdrawal-associated chromatin proteome was shared '
    'with or distinct from changes at other alcohol exposure time points, we constructed Venn '
    'diagrams comparing increased proteins across the three conditions (Fig. 3C). Of the 433 '
    'proteins increased in the chromatin fraction during acute withdrawal, 189 were uniquely '
    'increased at this time point and not significantly regulated during intoxication or '
    'protracted abstinence, indicating a withdrawal-specific recruitment of these proteins to '
    'chromatin. A substantial proportion — 152 proteins — were shared across all three '
    'conditions, suggesting that a core set of chromatin-associated proteins becomes persistently '
    'elevated beginning at intoxication and maintained throughout abstinence. An additional 83 '
    'proteins were shared between acute withdrawal and intoxication only, consistent with changes '
    'that are initiated during active alcohol exposure and amplified or sustained through '
    'withdrawal before resolving during protracted abstinence.'
)

body(doc,
    'Analysis of decreased chromatin proteins revealed an even more striking degree of '
    'withdrawal specificity (Fig. 3D). Of the 640 proteins decreased in the chromatin fraction '
    'during acute withdrawal, 491 — representing 77% of the total — were not significantly '
    'decreased during either intoxication or protracted abstinence, indicating that this wave of '
    'chromatin protein depletion is largely unique to the withdrawal state. By contrast, the '
    'number of proteins decreased exclusively during intoxication (n=27) or protracted '
    'abstinence (n=55) was substantially smaller. Only 42 proteins were commonly decreased '
    'across all three conditions, suggesting that shared persistent downregulation at chromatin '
    'is limited relative to the condition-specific responses. Together, these patterns indicate '
    'that acute withdrawal triggers a profound and predominantly state-specific reorganization '
    'of the chromatin-bound proteome in the amygdala, characterized by both a selective '
    'recruitment and a widespread depletion of chromatin-associated proteins.'
)

body(doc,
    '[PLACEHOLDER — GO network analysis of chromatin AW increased proteins: describe enriched '
    'biological processes, key proteins, and their relevance to chromatin regulation and '
    'alcohol-related neuroadaptation. Add figure panel reference (Fig. 3E or equivalent).]'
)

body(doc,
    '[PLACEHOLDER — GO network analysis of chromatin AW decreased proteins: describe enriched '
    'biological processes, key proteins, and interpretation. Add figure panel reference.]'
)

spacer(doc)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_DOCX)
print(f'Saved: {OUT_DOCX}')
