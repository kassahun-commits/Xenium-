from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = 'Heatmap_Methods.docx'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def heading(text, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def body(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Methods: Subcellular Compartment Heatmaps')
r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True

doc.add_paragraph()

# Overview
heading('Overview')
body(
    'Two sets of heatmaps were generated to visualise changes in protein abundance '
    'across three alcohol exposure conditions \u2014 Intoxication (I), Acute Withdrawal (AW), '
    'and Protracted Abstinence (PA) \u2014 relative to Na\u00efve (N) in four subcellular '
    'compartments of rat amygdala: Membrane, Cytosol, Chromatin, and Soluble Nuclear. '
    'Heatmap 1 includes proteins significant in at least one condition; '
    'Heatmap 2 (this folder) includes only proteins significant in all three conditions.'
)

doc.add_paragraph()

# Data source
heading('Data Source')
body(
    'Label-free quantification (LFQ) intensity values were extracted from the file '
    '\u201cEDIT Philipp Alcohol proteome Simplified \u2013 Jan 2024 copy.xlsx\u201d, '
    'using the individual compartment sheets: Membrane, Cytosol, Chromatin, and '
    'Soluble Nuclear. Each sheet contains LFQ values for 5 replicates per condition '
    '(a mix of female and male rats).'
)

doc.add_paragraph()

# Protein filtering
heading('Protein Filtering')
body('Proteins were retained as follows per compartment:')
bullet('Membrane: all proteins (no Filter column in this sheet)')
bullet('Cytosol: all proteins (no Filter column in this sheet)')
bullet('Chromatin: proteins labelled \u201cKeep\u201d or \u201cReview\u201d in the Filter column')
bullet(
    'Soluble Nuclear: all proteins regardless of Filter status '
    '(Keep, Review, and Exclude all included; n\u202f=\u202f3,013 total). '
    'This was done because the Soluble Nuclear data does not require the same '
    'quality filtering as Chromatin.'
)

doc.add_paragraph()

# FC and statistics
heading('Fold Change and Statistical Testing')
body(
    'For each protein and each condition, the mean LFQ intensity across the 5 '
    'condition replicates was subtracted from the mean LFQ intensity across the '
    '5 Na\u00efve replicates to produce a fold change (FC) value in log\u2082 LFQ '
    'space (FC\u202f=\u202fmean(condition)\u202f\u2212\u202fmean(Na\u00efve)).'
)
body(
    'Statistical significance was assessed using a two-sample Welch\u2019s t-test '
    '(unequal variance) between the Na\u00efve and condition replicates for each protein. '
    'P-values were corrected using a rank-based Benjamini\u2013Hochberg-style correction: '
    'corrected p = \u2212log\u2082(min(1, p \u00d7 N / rank)), where N is the total number '
    'of proteins with valid p-values and rank is the ascending rank of the p-value. '
    'Proteins were considered significant at corrected p\u202f>\u202f3.3 and '
    '|FC|\u202f>\u202f0.5.'
)

doc.add_paragraph()

# Protein selection for heatmap
heading('Protein Selection for Heatmap 2 (This Folder)')
body(
    'Only proteins that passed the significance threshold (corrected p > 3.3 AND '
    '|FC| > 0.5) in ALL THREE conditions simultaneously \u2014 Intoxication, Acute '
    'Withdrawal, AND Protracted Abstinence \u2014 were included. This conservative '
    'criterion identifies proteins that show consistent, statistically robust changes '
    'across the entire alcohol exposure timeline.'
)
body('Number of proteins meeting this criterion per compartment:')
bullet('Membrane: 19 proteins')
bullet('Cytosol: 11 proteins')
bullet('Chromatin: 201 proteins')
bullet('Soluble Nuclear: 18 proteins (all filter categories included)')

doc.add_paragraph()

# Heatmap construction
heading('Heatmap Construction')
body(
    'The FC values (log\u2082 LFQ, condition vs Na\u00efve) for each selected protein '
    'were arranged into a matrix with proteins as rows and the three conditions '
    '(Intoxication, Acute Withdrawal, Protracted Abstinence) as columns. '
    'Raw fold change values were used without z-scoring, so the colour scale '
    'directly reflects the magnitude of change versus Na\u00efve. Values were clipped '
    'at \u00b13 for visualisation purposes to prevent outliers from compressing '
    'the colour scale; the colourbar is labelled accordingly (\u2264\u22123 and \u2265+3).'
)
body(
    'Proteins (rows) were ordered by unsupervised hierarchical clustering using '
    'Ward linkage on Euclidean distances computed from the 3-column FC matrix. '
    'Missing values (proteins detected in some but not all conditions) were treated '
    'as 0 (no change) for the purpose of distance calculation only; the original NaN '
    'values were retained in the plotted matrix.'
)
body(
    'The colour scale uses a red\u2013white\u2013blue diverging palette: red indicates '
    'increased abundance versus Na\u00efve, blue indicates decreased abundance, and '
    'white indicates no change.'
)

doc.add_paragraph()

# Output
heading('Output Files')
bullet('Heatmap_SignificantAllConditions.pdf \u2014 4-page PDF, one page per compartment')
bullet('make_heatmap_allsig.py \u2014 Python script used to generate the heatmap')
bullet('Heatmap_Methods.docx \u2014 this document')

doc.save(OUT)
print(f'Saved: {OUT}')
