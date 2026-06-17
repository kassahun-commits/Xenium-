from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

doc = Document('Subcellular_Location_Annotation_Methods.docx')

def make_para(text, size=10.5, style_val='Normal'):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_val)
    pPr.append(pStyle)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2)))
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(int(size * 2)))
        rPr.append(sz); rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p

def make_bold_para(text, size=12):
    p = make_para(text, size=size)
    r = p.findall('.//' + qn('w:r'))[0]
    rPr = r.find(qn('w:rPr'))
    b = OxmlElement('w:b'); rPr.append(b)
    return p

def make_bullet(text, size=10.5):
    return make_para(text, size=size, style_val='ListBullet')

# Find the "Annotation performed..." paragraph
annot_para = None
for p in doc.paragraphs:
    if 'Annotation performed' in p.text:
        annot_para = p._element
        break

# Build new elements IN ORDER — each inserted before annot_para in sequence
new_elements = [
    make_para(''),
    make_bold_para('9. Rationale for Retaining Review Category in Primary Analysis', size=12),
    make_para(
        'The primary analysis uses proteins classified as Keep or Review (n\u2009=\u20092,192), '
        'with Exclude proteins (n\u2009=\u2009821) removed. This decision reflects the difficulty of '
        'distinguishing genuine nuclear presence from fractionation contamination using annotation '
        'data alone, particularly given the growing body of evidence for non-canonical nuclear '
        'localization of metabolic and mitochondrial proteins.'),
    make_para(''),
    make_bold_para('Why Review proteins are retained:', size=10.5),
    make_bullet(
        'Moonlighting and dual-localization are widespread. A systematic literature search of proteins '
        'in this dataset identified nuclear localization evidence for proteins with primary mitochondrial '
        'or metabolic annotations, including Gcdh (nuclear histone crotonylation in glioblastoma stem '
        'cells; Yuan et al., Nature 2023), Acat1 (nuclear NF-\u03baB regulation; Wei et al., 2025), '
        'Aldh2 (AMPK-induced nuclear translocation in rat kidney; Choi et al., 2011), and Acaa2 '
        '(nuclear thyroid hormone receptor coactivator; Wang & Ledee, 2021). Many of these nuclear '
        'roles were reported only recently, suggesting the literature is still catching up.'),
    make_bullet(
        'Alcohol exposure may itself induce novel nuclear translocation. Changes in nuclear protein '
        'composition are a plausible biological outcome of alcohol treatment; excluding proteins based '
        'on their canonical localization risks discarding condition-specific nuclear translocation '
        'events that could represent the primary finding of the study.'),
    make_bullet(
        'Unknown and Not found proteins cannot be classified as contaminants. Proteins lacking UniProt '
        'annotation (265 Unknown, 2 Not found) cannot be assigned to any compartment; their exclusion '
        'would introduce a systematic bias against unannotated or novel proteins.'),
    make_bullet(
        'Soluble mitochondrial proteins have no structural barrier to nuclear entry. Matrix enzymes '
        'involved in fatty acid oxidation, amino acid catabolism, and the TCA cycle lack transmembrane '
        'anchors and several have documented nuclear roles, making their presence in a chromatin '
        'fraction biologically plausible.'),
    make_para(''),
    make_bold_para('Why Exclude proteins are removed:', size=10.5),
    make_bullet(
        'Membrane-anchored and secreted proteins are structurally implausible nuclear residents. '
        'Proteins with confirmed transmembrane domains (plasma membrane, ER membrane, mitochondrial '
        'inner/outer membrane) or signal peptides directing secretion are unlikely to enter the '
        'nucleus. Their presence in a chromatin fraction most likely reflects carry-over of membrane '
        'vesicles or fragmented organelles during cell lysis and nuclear isolation.'),
    make_bullet(
        'Membrane contamination is a recognized artifact of chromatin fractionation. Incomplete '
        'removal of membrane fragments during nuclear isolation is a well-documented source of '
        'non-nuclear proteins in chromatin preparations, particularly for abundant membrane proteins.'),
    make_bullet(
        'Including membrane proteins inflates background noise and reduces statistical power to '
        'detect genuine chromatin-associated changes. Removing likely contaminants improves '
        'signal-to-noise ratio for the proteins of interest.'),
    make_para(''),
    make_para(
        'Results are presented in two filtered versions: Keep only (confirmed nuclear/cytoplasmic '
        'proteins, n\u2009=\u20091,736) and Keep\u2009+\u2009Review (all non-membrane proteins, '
        'n\u2009=\u20092,192). Concordance between the two versions provides a sensitivity check: '
        'findings that replicate across both filtered sets are more robust to annotation uncertainty. '
        'The full unfiltered dataset (all 3,013 proteins including Exclude) is retained for secondary '
        'analysis if needed.'),
    make_para(''),
]

for elem in new_elements:
    annot_para.addprevious(elem)

doc.save('Subcellular_Location_Annotation_Methods.docx')
print('Done. Final paragraphs:')
for i, p in enumerate(doc.paragraphs):
    if i >= 50:
        print(i, repr(p.text[:90]))
