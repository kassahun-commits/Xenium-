from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('Subcellular_Location_Annotation_Methods.docx')

def set_font(para, size=10.5, bold=False):
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.bold = bold

def add_para(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return p

# Find the index of the last paragraph ("Annotation performed...")
# We'll insert the new section before it
target_text = 'Annotation performed using Claude Code'
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if target_text in p.text:
        insert_idx = i
        break

# We'll rebuild by appending a new section heading + content before the final line.
# python-docx doesn't support true insertion, so we manipulate the XML directly.
from docx.oxml.ns import qn
from lxml import etree
import copy

body = doc.element.body

# Find the paragraph element for "Annotation performed..."
annot_para = doc.paragraphs[insert_idx]._element

def make_heading_para(doc, text, level_style='Heading 2'):
    p = doc.add_paragraph(style=level_style)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.bold = True
    return p._element

def make_normal_para(text, size=10.5):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p_elem = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p_elem.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_elem.append(r)
    return p_elem

def make_bullet_para(text, size=10.5):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    p_elem = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'ListBullet')
    pPr.append(pStyle)
    p_elem.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p_elem.append(r)
    return p_elem

new_elements = []

# Blank line
new_elements.append(make_normal_para(''))

# Section heading
new_elements.append(make_normal_para('9. Rationale for Retaining Review Category in Primary Analysis', size=12))

# Intro
new_elements.append(make_normal_para(
    'The primary analysis reported here uses proteins classified as Keep or Review (n = 2,192), '
    'with Exclude proteins (n = 821) removed. This decision reflects the difficulty of '
    'distinguishing genuine nuclear presence from fractionation contamination using annotation '
    'data alone, particularly given the growing body of evidence for non-canonical nuclear '
    'localization of metabolic and mitochondrial proteins.', size=10.5))

new_elements.append(make_normal_para(''))
new_elements.append(make_normal_para('Why Review proteins are retained:', size=10.5))

bullets_keep = [
    'Moonlighting and dual-localization are widespread. A systematic literature search of proteins '
    'in this dataset identified nuclear localization evidence for proteins with primary mitochondrial '
    'or metabolic annotations, including Gcdh (nuclear histone crotonylation in glioblastoma stem '
    'cells; Yuan et al., Nature 2023), Acat1 (nuclear NF-\u03baB regulation; Wei et al., 2025), '
    'Aldh2 (AMPK-induced nuclear translocation; Choi et al., 2011), and Acaa2 (nuclear thyroid '
    'hormone receptor coactivator; Wang & Ledee, 2021). Many of these nuclear roles were reported '
    'only recently, suggesting the literature is still catching up.',
    'Alcohol exposure may itself induce novel nuclear translocation. Changes in nuclear protein '
    'composition are a plausible biological outcome of alcohol treatment; excluding proteins based '
    'on their canonical localization risks discarding condition-specific nuclear translocation '
    'events that could represent the primary finding.',
    'Unknown and Not found proteins cannot be classified as contaminants. Proteins lacking UniProt '
    'annotation (265 Unknown, 2 Not found) cannot be assigned to any compartment; exclusion would '
    'introduce a systematic bias against unannotated or novel proteins.',
    'Mitochondrial proteins without membrane anchors are plausible nuclear residents. Soluble '
    'matrix proteins (e.g., TCA cycle enzymes, fatty acid oxidation enzymes) have no structural '
    'barrier preventing nuclear entry, and several have documented nuclear roles.',
]
for b in bullets_keep:
    new_elements.append(make_bullet_para(b))

new_elements.append(make_normal_para(''))
new_elements.append(make_normal_para('Why Exclude proteins are removed:', size=10.5))

bullets_excl = [
    'Membrane-anchored and secreted proteins are structurally implausible nuclear residents. '
    'Proteins with confirmed transmembrane domains (plasma membrane, ER membrane, mitochondrial '
    'inner/outer membrane) or signal peptides directing secretion are unlikely to enter the '
    'nucleus under normal or alcohol-exposed conditions. Their presence in a chromatin fraction '
    'most likely reflects carry-over of membrane vesicles during cell lysis.',
    'Membrane contamination is a recognized artifact of chromatin fractionation. Incomplete '
    'removal of membrane fragments during nuclear isolation is a well-documented source of '
    'non-nuclear proteins in chromatin preparations, particularly for abundant membrane proteins.',
    'Retaining membrane proteins would inflate background noise and reduce power to detect '
    'genuine chromatin-associated changes. The signal-to-noise ratio for detecting nuclear '
    'protein changes is improved by removing proteins that are almost certainly technical '
    'contaminants.',
]
for b in bullets_excl:
    new_elements.append(make_bullet_para(b))

new_elements.append(make_normal_para(''))
new_elements.append(make_normal_para(
    'Analyses are presented in two versions: Keep only (confirmed nuclear/cytoplasmic proteins, '
    'n = 1,736) and Keep + Review (all non-membrane proteins, n = 2,192). Comparing these two '
    'versions provides a sensitivity check: findings that replicate across both filtered sets '
    'are more likely to reflect genuine biology rather than annotation uncertainty. The full '
    'unfiltered dataset (all 3,013 proteins including Exclude) is retained and available for '
    'secondary analysis if needed.', size=10.5))

new_elements.append(make_normal_para(''))

# Insert all new elements before the "Annotation performed..." paragraph
for elem in reversed(new_elements):
    annot_para.addprevious(elem)

doc.save('Subcellular_Location_Annotation_Methods.docx')
print('Saved.')
