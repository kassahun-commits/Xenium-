"""
Generates Protein_Chromatin_Roles.docx
Table: Protein | Primary Function | Nuclear/Chromatin Role | Source
Landscape, clean formatting, easy to paste into PowerPoint.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "Protein_Chromatin_Roles.docx"

# ── Data ──────────────────────────────────────────────────────────────────────
HEADER_COLOR  = "1F3864"   # dark navy
ALT_ROW_COLOR = "EBF0FA"   # light blue-grey
WHITE         = "FFFFFF"
HEADER_TEXT   = "FFFFFF"

PROTEINS = [
    {
        "name":     "Adh5",
        "primary":  "Detoxifies formaldehyde via glutathione (formaldehyde dehydrogenase / GSNOR); metabolizes S-nitrosoglutathione; alcohol metabolism",
        "nuclear":  "Histone and DNA demethylation reactions (by KDM and TET enzymes) release formaldehyde as a byproduct inside the nucleus. Adh5 scavenges this nuclear formaldehyde to prevent DNA crosslinks — directly coupled to chromatin remodeling and epigenetic regulation",
        "source":   "Pontel et al. (2015) Mol Cell 60:177–188\nSanghvi et al. (2021) Science Advances 7:eabd7197",
    },
    {
        "name":     "Asna1",
        "primary":  "Delivers tail-anchored (TA) proteins to the ER membrane (TRC40/GET3 pathway); chaperones unfolded proteins under oxidative stress and arsenite exposure",
        "nuclear":  "Detected in nucleoli and perinuclear compartments; nucleolar signal co-localises with known nucleolar markers; may support nuclear membrane protein insertion and stress responses — specific chromatin function not yet defined",
        "source":   "Stefanovic & Hegde (2007) Cell 128:1147–1159\nHuman Protein Atlas (proteinatlas.org/ENSG00000198356)",
    },
    {
        "name":     "Cryz",
        "primary":  "NADPH-dependent quinone reductase (ζ-crystallin); detoxifies reactive quinones; also acts as a structural lens protein",
        "nuclear":  "Nuclear function not well characterised; some crystallin family members associate with chromatin as chaperones and have been detected in the nucleus under stress conditions; specific chromatin role in amygdala not established",
        "source":   "Rao et al. (1997) J Biol Chem 272:2529–2535\nUniProt Q08257",
    },
    {
        "name":     "Dhrs4",
        "primary":  "NADPH-dependent short-chain dehydrogenase/reductase; converts retinol → retinal in retinoic acid biosynthesis; reduces carbonyl compounds",
        "nuclear":  "Modulates local retinoic acid (RA) availability; RA activates RAR/RXR nuclear receptors that bind response elements on chromatin to regulate transcription of target genes — indirect but important chromatin regulatory axis",
        "source":   "Matsunaga et al. (2008) Drug Metab Pharmacokinet 23:207–215\nMano et al. (2000) Biochemistry 39:15327–15334",
    },
    {
        "name":     "Echdc3",
        "primary":  "Enoyl-CoA hydratase domain protein; involved in fatty acid β-oxidation intermediate processing; primarily mitochondrial",
        "nuclear":  "Direct nuclear/chromatin function not characterised; fatty acid metabolism enzymes in the nucleus have been linked to local acetyl-CoA supply for histone acetylation, but this has not been confirmed for Echdc3 specifically",
        "source":   "UniProt Q5SXR6\nKoronowski & Bhatt (2022) Mol Metab 65:101582 (general metabolic enzymes in nucleus review)",
    },
    {
        "name":     "Eif3f",
        "primary":  "Core subunit of the 13-subunit eIF3 translation initiation complex; scaffolds ribosome 43S pre-initiation complex assembly; links mTORC1–S6K1 anabolic signalling; regulates muscle mass via MAFbx/Atrogin-1 axis",
        "nuclear":  "A distinct nuclear eIF3 complex exists and shuttles between nucleus and cytoplasm in a cell-cycle-dependent manner; Eif3f deubiquitinates activated NOTCH1 to promote nuclear import; nuclear eIF3 may couple transcription to translation initiation co-transcriptionally",
        "source":   "Csibi et al. (2010) Sci Signal 3:ra29\nProg & Etienne-Manneville (2014) Cell Mol Life Sci 71:3257–3267\nGeneCards EIF3F",
    },
    {
        "name":     "Gckr",
        "primary":  "Glucokinase regulatory protein; binds and inhibits glucokinase; modulates glucose-sensing and glycolytic flux; expression in liver and brain (tanycytes, neurons)",
        "nuclear":  "Predominantly nuclear — sequesters glucokinase in the nucleus at low glucose; releases it to cytoplasm when glucose rises (fructose metabolites modulate affinity). Acts as a nuclear glucose sensor, directly linking metabolic state to nuclear signalling",
        "source":   "Agius (2008) Biochem J 414:1–18\nGonzalez et al. (2015) Int J Mol Sci 16:7377–7393\nRivero-Gutiérrez et al. (2019) Front Neurosci 13:275",
    },
    {
        "name":     "Glyat",
        "primary":  "Glycine N-acyltransferase; conjugates glycine to acyl-CoA compounds for amino acid detoxification (e.g., benzoyl-CoA → hippuric acid); amino acid metabolism",
        "nuclear":  "No established nuclear or chromatin role; acyl-CoA-handling enzymes in the nucleus can potentially influence histone acylation marks, but no direct evidence exists for Glyat",
        "source":   "Webster et al. (1976) Biochem J 160:535–544\nUniProt Q8CHR6",
    },
    {
        "name":     "Idh3a",
        "primary":  "Catalytic α-subunit of mitochondrial NAD⁺-dependent isocitrate dehydrogenase (IDH3); catalyses the rate-limiting TCA cycle step: isocitrate → α-ketoglutarate (αKG) + CO₂ + NADH",
        "nuclear":  "αKG produced by IDH enzymes is an essential co-factor for TET family DNA demethylases (5mC → 5hmC) and KDM/Jumonji-C histone demethylases; nuclear Idh3a could locally supply αKG to directly fuel epigenetic modifications on chromatin",
        "source":   "Xu et al. (2011) Science 332:1359–1364\nYe et al. (2018) Science Advances 4:eaaw4543\nUniProt P56471",
    },
    {
        "name":     "Mmut",
        "primary":  "Methylmalonyl-CoA mutase; converts methylmalonyl-CoA → succinyl-CoA (requires adenosylcobalamin / vitamin B12); catabolism of branched-chain amino acids, odd-chain fatty acids, and cholesterol",
        "nuclear":  "Succinyl-CoA (the product of Mmut) is a substrate for histone succinylation, a recently described post-translational modification catalysed by KAT2A/GCN5. Nuclear Mmut could locally supply succinyl-CoA for histone modification, but direct evidence in brain is limited",
        "source":   "Weinert et al. (2013) Nat Chem Biol 9:833–838 (histone succinylation)\nUniProt P27653",
    },
    {
        "name":     "Prps1",
        "primary":  "Phosphoribosyl pyrophosphate (PRPP) synthetase 1; synthesises PRPP from ribose-5-phosphate + ATP; rate-limiting step for de novo and salvage synthesis of purines and pyrimidines (DNA/RNA building blocks)",
        "nuclear":  "Nuclear presence likely supports chromatin-proximal nucleotide synthesis for DNA replication and repair; PRPP is also a precursor for NAD⁺ biosynthesis (via NMN), linking Prps1 to the supply of NAD⁺ for sirtuin histone deacetylases (HDAC class III) and PARP-mediated DNA repair",
        "source":   "Hove-Jensen et al. (2017) Microbiol Mol Biol Rev 81:e00040-16\nMedlinePlus Genetics: PRPS1\nCanto et al. (2015) Cell Metab 22:31–53 (NAD⁺–sirtuin axis)",
    },
]

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Set landscape, US Letter
section = doc.sections[0]
section.orientation    = WD_ORIENT.LANDSCAPE
section.page_width     = Inches(11)
section.page_height    = Inches(8.5)
section.left_margin    = Inches(0.6)
section.right_margin   = Inches(0.6)
section.top_margin     = Inches(0.6)
section.bottom_margin  = Inches(0.6)

# Title
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("Chromatin-Fraction Proteins — Rat Amygdala: Normal Roles and Nuclear/Chromatin Functions")
title_run.bold      = True
title_run.font.size = Pt(14)
title_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()  # small spacer

# ── Helper: set cell shading ──────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """kwargs: top, bottom, left, right — each a dict with color, sz, val"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        if side in kwargs:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'),   kwargs[side].get('val',   'single'))
            b.set(qn('w:sz'),    str(kwargs[side].get('sz', 4)))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), kwargs[side].get('color', 'auto'))
            tcBorders.append(b)
    tcPr.append(tcBorders)

# ── Table ─────────────────────────────────────────────────────────────────────
# Content width = 11 - 1.2 = 9.8 inches
# Cols: Protein(0.9) | Primary Function(2.5) | Nuclear/Chromatin Role(3.8) | Source(2.6)
COL_W = [Inches(0.9), Inches(2.5), Inches(3.8), Inches(2.6)]
HEADERS = ["Protein", "Primary Cellular Function", "Nuclear / Chromatin Role", "Source"]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Set column widths
for i, cell in enumerate(table.columns[i].cells[0] for i in range(4)):
    pass  # we'll set widths on each cell

# Header row
hdr_cells = table.rows[0].cells
for i, (cell, hdr) in enumerate(zip(hdr_cells, HEADERS)):
    shade_cell(cell, HEADER_COLOR)
    cell.width = COL_W[i]
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(hdr)
    run.bold           = True
    run.font.size      = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Data rows
for ri, prot in enumerate(PROTEINS):
    row_cells = table.add_row().cells
    bg = ALT_ROW_COLOR if ri % 2 == 0 else WHITE

    data = [prot["name"], prot["primary"], prot["nuclear"], prot["source"]]
    for ci, (cell, text) in enumerate(zip(row_cells, data)):
        shade_cell(cell, bg)
        cell.width = COL_W[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Handle multi-line source text
        lines = text.split('\n')
        for li, line in enumerate(lines):
            if li == 0:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(line)
            run.font.size = Pt(8.5) if ci == 0 else Pt(8.5)
            if ci == 0:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

# Fix table column widths explicitly via XML (python-docx sometimes ignores Inches on cells)
from docx.oxml.ns import nsmap
tbl = table._tbl
tblPr = tbl.tblPr
tblW = OxmlElement('w:tblW')
tblW.set(qn('w:w'),    '9072')   # 9.8 * 1440 / 1.08 ≈ total in twips... use simpler approach
tblW.set(qn('w:type'), 'dxa')

# Set each cell width properly
from docx.oxml import OxmlElement as OE
width_twips = [int(w.inches * 1440) for w in COL_W]   # 1 inch = 1440 twips

for row in table.rows:
    for ci, cell in enumerate(row.cells):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove any existing tcW
        for old in tcPr.findall(qn('w:tcW')):
            tcPr.remove(old)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(width_twips[ci]))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.insert(0, tcW)

# ── Footer note ────────────────────────────────────────────────────────────────
doc.add_paragraph()
note_para = doc.add_paragraph()
note_run = note_para.add_run(
    "Note: These proteins were identified in the chromatin fraction of rat amygdala proteomics data. "
    "Many are metabolic enzymes with emerging 'moonlighting' roles in the nucleus, locally supplying "
    "metabolites (α-KG, formaldehyde, succinyl-CoA, PRPP/NAD⁺) that serve as substrates or co-factors "
    "for chromatin-modifying enzymes (TET, KDM, sirtuins, PARP)."
)
note_run.italic    = True
note_run.font.size = Pt(8)
note_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(OUT)
print(f"Saved: {OUT}")
